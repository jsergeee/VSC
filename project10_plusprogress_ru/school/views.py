from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.db.models import Sum, Count, Q, Prefetch, Avg
from django.utils import timezone
from django.http import HttpResponse, JsonResponse
from django.views.decorators.http import require_GET, require_POST
from django.urls import reverse
from django.template.loader import render_to_string
from django.db import connection, transaction
from decimal import Decimal
from datetime import datetime, date, timedelta
from .telegram import notify_payment
import json
import csv
import uuid
import io
import os
import tempfile
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from weasyprint import HTML
import logging
from .models import Feedback
import traceback
from django.utils import timezone
from .utils import log_user_action
from .forms import TelegramSettingsForm
from django.conf import settings
import json
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import get_user_model
import requests
from django.conf import settings
from rest_framework import generics 
from .serializers import RegisterSerializer
from rest_framework import viewsets, permissions, status, generics
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.authtoken.models import Token
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator  
from .permissions import IsTeacherUser, IsStudentUser
from django.http import JsonResponse
from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import render, redirect
from django.contrib import messages
import openpyxl
from datetime import datetime
import traceback



# Импорты моделей
from .models import (
    User, Teacher, Student, Subject, LessonFormat, Lesson,
    LessonAttendance, LessonReport, Payment, Schedule, TrialRequest,
    Material, Deposit, StudentNote, GroupLesson, GroupEnrollment,
    Notification, LessonFeedback, TeacherRating, Homework,
    HomeworkSubmission, ScheduleTemplate, ScheduleTemplateStudent,
    StudentSubjectPrice, EmailVerificationToken, PaymentRequest
)

from .forms import (
    UserRegistrationForm, UserLoginForm, TrialRequestForm,
    LessonReportForm, ProfileUpdateForm, LessonFeedbackForm,
    HomeworkForm, HomeworkSubmissionForm, HomeworkCheckForm,
    ScheduleTemplateForm, PublicFeedbackForm
)

from .utils import send_verification_email, send_verification_success_email

logger = logging.getLogger(__name__)


# ============================================
# ЧАСТЬ 1: HELPER-КЛАССЫ ДЛЯ ФИНАНСОВЫХ РАСЧЕТОВ
# ============================================

class LessonFinanceCalculator:
    """
    ЕДИНЫЙ КАЛЬКУЛЯТОР ФИНАНСОВ ДЛЯ УРОКА
    """

    def __init__(self, lesson):
        self.lesson = lesson
        self.attendances = lesson.attendance.all()

    @property
    def total_cost(self) -> Decimal:
        """Общая стоимость урока для всех учеников"""
        return sum((a.cost for a in self.attendances), Decimal('0'))

    @property
    def teacher_payment(self) -> Decimal:
        """Общая выплата учителю за урок"""
        return sum((a.teacher_payment_share for a in self.attendances), Decimal('0'))

    @property
    def attended_cost(self) -> Decimal:
        """Стоимость только для присутствовавших"""
        return sum((a.cost for a in self.attendances if a.status == 'attended'), Decimal('0'))

    @property
    def attended_payment(self) -> Decimal:
        """Выплата учителю только за присутствовавших"""
        return sum((a.teacher_payment_share for a in self.attendances if a.status == 'attended'), Decimal('0'))

    @property
    def debt_cost(self) -> Decimal:
        """Стоимость уроков в долг"""
        return sum((a.cost for a in self.attendances if a.status == 'debt'), Decimal('0'))

    @property
    def stats(self) -> dict:
        """Полная статистика по уроку"""
        return {
            # Денежные показатели
            'total_cost': float(self.total_cost),
            'teacher_payment': float(self.teacher_payment),
            'attended_cost': float(self.attended_cost),
            'attended_payment': float(self.attended_payment),
            'debt_cost': float(self.debt_cost),

            # Количественные показатели
            'students_total': self.attendances.count(),
            'students_attended': self.attendances.filter(status='attended').count(),
            'students_debt': self.attendances.filter(status='debt').count(),
            'students_absent': self.attendances.filter(status='absent').count(),
            'students_registered': self.attendances.filter(status='registered').count(),
        }

    def get_attendance_details(self) -> list:
        """Детализация по ученикам"""
        return [
            {
                'id': a.id,
                'student_id': a.student.id,
                'student_name': a.student.user.get_full_name(),
                'cost': float(a.cost),
                'teacher_payment': float(a.teacher_payment_share),
                'status': a.status,
                # УБРАЛИ balance_before и balance_after
            }
            for a in self.attendances
        ]


class PeriodFinanceCalculator:
    """
    КАЛЬКУЛЯТОР ФИНАНСОВ ЗА ПЕРИОД
    Использовать для отчетов и дашбордов
    """

    def __init__(self, lessons_queryset, payments_queryset=None):
        self.lessons = lessons_queryset
        self.payments = payments_queryset if payments_queryset is not None else Payment.objects.none()

    @property
    def lessons_stats(self) -> dict:
        """Статистика по урокам за период"""
        completed = self.lessons.filter(status='completed')

        total_cost = Decimal('0')
        total_payment = Decimal('0')

        for lesson in completed:
            calc = LessonFinanceCalculator(lesson)
            total_cost += calc.total_cost
            total_payment += calc.teacher_payment

        return {
            'total': self.lessons.count(),
            'completed': completed.count(),
            'cancelled': self.lessons.filter(status='cancelled').count(),
            'overdue': self.lessons.filter(status='overdue').count(),
            'scheduled': self.lessons.filter(status='scheduled').count(),
            'total_cost': float(total_cost),
            'teacher_payment': float(total_payment),
        }

    @property
    def payments_stats(self) -> dict:
        """Статистика по платежам за период"""
        return {
            'income': float(self.payments.filter(payment_type='income').aggregate(Sum('amount'))['amount__sum'] or 0),
            'expense': float(self.payments.filter(payment_type='expense').aggregate(Sum('amount'))['amount__sum'] or 0),
            'teacher_payments': float(
                self.payments.filter(payment_type='teacher_payment').aggregate(Sum('amount'))['amount__sum'] or 0),
            'total': float(self.payments.aggregate(Sum('amount'))['amount__sum'] or 0),
            'count': self.payments.count(),
        }

    @property
    def school_finances(self) -> dict:
        """Финансовые показатели школы"""
        payments = self.payments_stats

        return {
            'income': payments['expense'],
            'expense': payments['teacher_payments'],
            'profit': payments['expense'] - payments['teacher_payments'],
            'profit_margin': ((payments['expense'] - payments['teacher_payments']) / payments['expense'] * 100) if
            payments['expense'] > 0 else 0
        }


class StudentFinanceHelper:
    """
    ПОМОЩНИК ПО СТАТИСТИКЕ УЧЕНИКА
    """

    def __init__(self, student):
        self.student = student
        self.user = student.user

    # УДАЛЯЕМ методы balance и debt

    def get_lessons_stats(self, days=30) -> dict:
        """Статистика по урокам за последние N дней"""
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days)

        attendances = LessonAttendance.objects.filter(
            student=self.student,
            lesson__date__gte=start_date,
            lesson__date__lte=end_date
        )

        total_cost = attendances.aggregate(Sum('cost'))['cost__sum'] or 0

        return {
            'period_days': days,
            'total': attendances.count(),
            'attended': attendances.filter(status='attended').count(),
            'debt': attendances.filter(status='debt').count(),
            'total_cost': float(total_cost),
            'average_cost': float(total_cost / attendances.count()) if attendances.exists() else 0
        }

    def get_lessons_stats_by_period(self, start_date=None, end_date=None):
        """Статистика по урокам за указанный период"""
        attendances = LessonAttendance.objects.filter(
            student=self.student
        )

        if start_date:
            attendances = attendances.filter(lesson__date__gte=start_date)
        if end_date:
            attendances = attendances.filter(lesson__date__lte=end_date)

        total_cost = attendances.aggregate(Sum('cost'))['cost__sum'] or 0

        return {
            'total': attendances.count(),
            'attended': attendances.filter(status='attended').count(),
            'debt': attendances.filter(status='debt').count(),
            'total_cost': float(total_cost),
            'average_cost': float(total_cost / attendances.count()) if attendances.exists() else 0
        }


class TeacherFinanceHelper:
    """
    ПОМОЩНИК ПО ФИНАНСАМ УЧИТЕЛЯ
    """

    def __init__(self, teacher):
        self.teacher = teacher
        self.user = teacher.user

    @property
    def wallet_balance(self) -> Decimal:
        """Текущий баланс кошелька"""
        return self.teacher.wallet_balance

    def get_payment_stats(self, days=30) -> dict:
        """Статистика выплат"""
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days)

        payments = Payment.objects.filter(
            user=self.user,
            payment_type='teacher_payment',
            created_at__date__gte=start_date,
            created_at__date__lte=end_date
        )

        return {
            'period_days': days,
            'total': payments.aggregate(Sum('amount'))['amount__sum'] or 0,
            'count': payments.count(),
            'average': (payments.aggregate(Sum('amount'))[
                            'amount__sum'] or 0) / payments.count() if payments.exists() else 0
        }


# ============================================
# ЧАСТЬ 2: ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ПОИСКА
# ============================================

def find_teacher_by_full_name(name):
    """Поиск учителя по полному имени"""
    if not name:
        return None

    name = str(name).strip()
    if not name:
        return None

    name_parts = name.split()
    if not name_parts:
        return None

    last_name = name_parts[0]
    teachers = Teacher.objects.filter(user__last_name__icontains=last_name)

    if teachers.exists():
        if teachers.count() == 1:
            return teachers.first()

        if len(name_parts) > 1:
            first_name = name_parts[1]
            teachers = teachers.filter(user__first_name__icontains=first_name)
            if teachers.exists():
                return teachers.first()

    for teacher in Teacher.objects.all():
        full_name = teacher.user.get_full_name().lower()
        if name.lower() in full_name:
            return teacher

    return None


def find_student_by_full_name(name):
    """Поиск ученика по полному имени"""
    if not name:
        return None

    name = str(name).strip()
    if not name:
        return None

    name_parts = name.split()
    if not name_parts:
        return None

    last_name = name_parts[0]
    students = Student.objects.filter(user__last_name__icontains=last_name)

    if students.exists():
        if students.count() == 1:
            return students.first()

        if len(name_parts) > 1:
            first_name = name_parts[1]
            students = students.filter(user__first_name__icontains=first_name)
            if students.exists():
                return students.first()

    for student in Student.objects.all():
        full_name = student.user.get_full_name().lower()
        if name.lower() in full_name:
            return student

    return None


def find_teacher_by_id(teacher_id):
    """Поиск учителя по ID"""
    try:
        return Teacher.objects.get(id=int(teacher_id))
    except (ValueError, Teacher.DoesNotExist):
        return None


def find_student_by_id(student_id):
    """Поиск ученика по ID"""
    try:
        return Student.objects.get(id=int(student_id))
    except (ValueError, Student.DoesNotExist):
        return None


def create_lesson_with_prices(teacher, student, subject, date, start_time, end_time):
    """Создание урока с автоматической подстановкой цен"""
    
    # Теперь передаем teacher в метод get_price_for
    cost, teacher_payment = StudentSubjectPrice.get_price_for(student, subject, teacher)

    if cost is None:
        cost = Decimal('1000')
    if teacher_payment is None:
        teacher_payment = cost * Decimal('0.7')

    lesson = Lesson.objects.create(
        teacher=teacher,
        subject=subject,
        date=date,
        start_time=start_time,
        end_time=end_time,
        base_cost=cost,
        base_teacher_payment=teacher_payment
    )

    LessonAttendance.objects.create(
        lesson=lesson,
        student=student,
        cost=cost,
        teacher_payment_share=teacher_payment,
        status='registered'
    )

    return lesson
# ============================================
# ЧАСТЬ 3: ОСНОВНЫЕ VIEWS (рефакторинг ключевых функций)
# ============================================

def home(request):
    """
    Главная страница
    """
    trial_form = TrialRequestForm()
    feedback_form = PublicFeedbackForm()

    # Получаем активные отзывы для главной
    feedbacks = Feedback.objects.filter(
        is_active=True,
        is_on_main=True
    ).select_related('teacher__user').order_by('sort_order', '-created_at')

    # ===== НОВАЯ СТАТИСТИКА ПО ОТЗЫВАМ =====
    # Все активные отзывы (не только на главной)
    all_active_feedbacks = Feedback.objects.filter(is_active=True)
    
    total_feedbacks = all_active_feedbacks.count()
    
    # Статистика по оценкам
    rating_stats = {
        'avg': all_active_feedbacks.aggregate(Avg('rating'))['rating__avg'] or 0,
        '5': all_active_feedbacks.filter(rating=5).count(),
        '4': all_active_feedbacks.filter(rating=4).count(),
        '3': all_active_feedbacks.filter(rating=3).count(),
        '2': all_active_feedbacks.filter(rating=2).count(),
        '1': all_active_feedbacks.filter(rating=1).count(),
    }
    
    # Проценты для визуализации
    if total_feedbacks > 0:
        rating_stats['5_percent'] = round(rating_stats['5'] / total_feedbacks * 100)
        rating_stats['4_percent'] = round(rating_stats['4'] / total_feedbacks * 100)
        rating_stats['3_percent'] = round(rating_stats['3'] / total_feedbacks * 100)
        rating_stats['2_percent'] = round(rating_stats['2'] / total_feedbacks * 100)
        rating_stats['1_percent'] = round(rating_stats['1'] / total_feedbacks * 100)
    else:
        rating_stats['5_percent'] = rating_stats['4_percent'] = rating_stats['3_percent'] = rating_stats['2_percent'] = rating_stats['1_percent'] = 0

    # Последние 3 отзыва для превью (если нужно)
    recent_feedbacks = all_active_feedbacks.order_by('-created_at')[:3]
    # ===== КОНЕЦ НОВОЙ СТАТИСТИКИ =====

    if request.method == 'POST':
        if 'trial_form' in request.POST:
            trial_form = TrialRequestForm(request.POST)
            if trial_form.is_valid():
                trial_form.save()
                messages.success(request, 'Заявка успешно отправлена! Мы свяжемся с вами в ближайшее время.')
                return redirect('home')
        elif 'feedback_form' in request.POST:
            feedback_form = PublicFeedbackForm(request.POST)
            if feedback_form.is_valid():
                feedback = feedback_form.save(commit=False)
                feedback.is_active = False
                feedback.is_on_main = False
                feedback.save()

                # Уведомление админу
                admin_users = User.objects.filter(is_superuser=True)
                for admin in admin_users:
                    Notification.objects.create(
                        user=admin,
                        title='✍️ Новый отзыв',
                        message=f'Новый отзыв от {feedback.name} ожидает модерации',
                        notification_type='system',
                        link=f'/admin/school/feedback/{feedback.id}/change/'
                    )

                messages.success(
                    request,
                    'Спасибо за ваш отзыв! Он появится на сайте после проверки.'
                )
                return redirect('home')

    context = {
        'trial_form': trial_form,
        'feedback_form': feedback_form,
        'subjects': Subject.objects.all(),
        'feedbacks': feedbacks,
        # ===== ДОБАВЛЯЕМ В КОНТЕКСТ =====
        'total_feedbacks': total_feedbacks,
        'rating_stats': rating_stats,
        'recent_feedbacks': recent_feedbacks,
        # ===== КОНЕЦ ДОБАВЛЕНИЯ =====
    }
    return render(request, 'school/home.html', context)


def register(request):
    """Регистрация пользователя"""
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                user = form.save()

                if send_verification_email(user, request):
                    messages.success(
                        request,
                        'Регистрация прошла успешно! На ваш email отправлено письмо с подтверждением.'
                    )
                else:
                    messages.warning(
                        request,
                        'Регистрация прошла успешно, но не удалось отправить письмо подтверждения.'
                    )

                return redirect('login')

            except Exception as e:
                messages.error(request, f'Ошибка при регистрации: {str(e)}')
        else:
            messages.error(request, 'Пожалуйста, исправьте ошибки в форме')
    else:
        form = UserRegistrationForm()

    return render(request, 'school/register.html', {'form': form})


def user_login(request):
    """Вход в систему"""
    if request.method == 'POST':
        form = UserLoginForm(request.POST)

        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']

            user = authenticate(username=username, password=password)

            if user:
                if not user.is_email_verified:
                    messages.warning(
                        request,
                        'Пожалуйста, подтвердите ваш email перед входом в систему. '
                        '<a href="{}" class="alert-link">Отправить письмо повторно</a>'.format(
                            reverse('resend_verification')
                        )
                    )
                    return redirect('login')

                login(request, user)

                # ✅ ЛОГИРОВАНИЕ ВХОДА
                from .utils import log_user_action
                log_user_action(
                    request,
                    'login',
                    f'Вход в систему (роль: {user.get_role_display()})',
                    additional_data={'role': user.role}
                )
                # ✅ КОНЕЦ ЛОГИРОВАНИЯ

                if user.role == 'student':
                    return redirect('student_dashboard')
                elif user.role == 'teacher':
                    return redirect('teacher_dashboard')
                else:
                    return redirect('admin:index')
            else:
                messages.error(request, 'Неверное имя пользователя или пароль')
    else:
        form = UserLoginForm()

    return render(request, 'school/login.html', {'form': form})


@login_required
def user_logout(request):
    """Выход из системы"""
    log_user_action(request, 'logout', 'Выход из системы')
    logout(request)
    return redirect('home')


@login_required
def dashboard(request):
    """Редирект на соответствующий дашборд"""
    user = request.user

    if user.role == 'student':
        return redirect('student_dashboard')
    elif user.role == 'teacher':
        return redirect('teacher_dashboard')
    else:
        return redirect('admin:index')


@login_required
def student_dashboard(request):
    """Личный кабинет ученика"""

    # ===== ПРОВЕРКА ДОСТУПА =====
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    user = request.user

    # ===== ПРОФИЛЬ УЧЕНИКА =====
    try:
        student = user.student_profile
    except:
        student = Student.objects.create(user=user)
        messages.info(request, 'Профиль ученика был создан')

    # ===== ФИНАНСЫ: ИСПОЛЬЗУЕМ НОВЫЙ МЕТОД =====
    from django.db.models import Sum
    from school.models import Payment, LessonAttendance

    # Пополнения (income)
    total_deposits = Payment.objects.filter(
        user=user,
        payment_type='income'
    ).aggregate(Sum('amount'))['amount__sum'] or 0

    # Списания за проведенные занятия (expense)
    total_expenses = Payment.objects.filter(
        user=user,
        payment_type='expense'
    ).aggregate(Sum('amount'))['amount__sum'] or 0

    # Стоимость проведенных уроков (attended)
    attended_cost = LessonAttendance.objects.filter(
        student=student,
        status='attended'
    ).aggregate(Sum('cost'))['cost__sum'] or 0

    # ✅ БАЛАНС из нового метода
    current_balance = user.get_balance()

    # Для отладки
    print(f"\n{'💰' * 30}")
    print(f"💰 Баланс ученика {user.username}:")
    print(f"   Пополнения: {total_deposits}")
    print(f"   Списания (expense): {total_expenses}")
    print(f"   Проведено уроков на: {attended_cost}")
    print(f"   Баланс (get_balance): {current_balance}")
    print(f"{'💰' * 30}\n")

    # ===== СТАТИСТИКА ПО УРОКАМ =====
    attended_lessons = LessonAttendance.objects.filter(
        student=student,
        status='attended'
    ).count()

    attended_cost_total = LessonAttendance.objects.filter(
        student=student,
        status='attended'
    ).aggregate(Sum('cost'))['cost__sum'] or 0

    debt_lessons = LessonAttendance.objects.filter(
        student=student,
        status='debt'
    ).count()

    debt_cost_total = LessonAttendance.objects.filter(
        student=student,
        status='debt'
    ).aggregate(Sum('cost'))['cost__sum'] or 0

    # ===== УЧИТЕЛЯ И ПОПОЛНЕНИЯ =====
    teachers = student.teachers.all()
    recent_deposits = student.deposits.all()[:5]

    # ===== УРОКИ ДЛЯ КАЛЕНДАРЯ =====
    all_lessons = Lesson.objects.filter(
        attendance__student=student
    ).select_related('teacher__user', 'subject', 'format').distinct().order_by('date', 'start_time')

    # ===== ПРОВЕРКА ПРОСРОЧЕННЫХ УРОКОВ =====
    from datetime import datetime
    updated_count = 0
    for lesson in all_lessons:
        if lesson.status == 'scheduled':
            lesson_datetime = datetime.combine(lesson.date, lesson.start_time)
            if lesson_datetime < datetime.now():
                lesson.status = 'overdue'
                lesson.save()
                updated_count += 1
                print(f"⚠️ Урок {lesson.id} автоматически помечен как просроченный (из дашборда)")

    if updated_count > 0:
        print(f"✅ Обновлено {updated_count} просроченных уроков")
        # Обновляем queryset после изменений
        all_lessons = Lesson.objects.filter(
            attendance__student=student
        ).select_related('teacher__user', 'subject', 'format').distinct().order_by('date', 'start_time')

    # ===== БЛИЖАЙШИЕ УРОКИ (ДЛЯ СПИСКА) =====
    upcoming_lessons_list = Lesson.objects.filter(
        attendance__student=student,
        date__gte=date.today(),
        status='scheduled'
    ).select_related('teacher__user', 'subject', 'format').distinct().order_by('date', 'start_time')[:10]

    # ===== ПРОШЕДШИЕ УРОКИ =====
    past_lessons = Lesson.objects.filter(
        attendance__student=student,
        status='completed'
    ).select_related('teacher__user', 'subject').distinct().order_by('-date')[:10]

    # ===== МЕТОДИЧЕСКИЕ МАТЕРИАЛЫ =====
    from school.models import Material
    from django.db.models import Q

    materials = Material.objects.filter(
        Q(students=student) | Q(is_public=True) | Q(teachers__in=teachers)
    ).distinct()[:20]

    # ===== ДОМАШНИЕ ЗАДАНИЯ =====
    from school.models import Homework

    recent_homeworks = Homework.objects.filter(
        student=student,
        is_active=True
    ).exclude(
        submission__status='checked'
    ).select_related('teacher__user', 'subject').order_by('deadline')[:4]

    # Статистика по домашним заданиям
    all_homeworks = Homework.objects.filter(student=student)
    homework_stats = {
        'total': all_homeworks.count(),
        'pending': sum(1 for h in all_homeworks if h.get_status() == 'pending'),
        'submitted': sum(1 for h in all_homeworks if h.get_status() == 'submitted'),
        'checked': sum(1 for h in all_homeworks if h.get_status() == 'checked'),
        'overdue': sum(1 for h in all_homeworks if h.get_status() == 'overdue'),
    }

    # ===== ГРУППОВЫЕ УРОКИ =====
    from school.models import GroupLesson

    group_lessons = GroupLesson.objects.filter(
        enrollments__student=student
    ).select_related('teacher__user', 'subject')

    # ===== КАЛЕНДАРЬ =====
    calendar_events = []

    # Цвета для разных статусов
    status_colors = {
        'scheduled': '#007bff',  # синий
        'completed': '#28a745',  # зеленый
        'cancelled': '#dc3545',  # красный
        'overdue': '#fd7e14',  # оранжевый
        'rescheduled': '#ffc107',  # желтый
        'no_show': '#6c757d',  # серый
    }

    # Обычные уроки
    for lesson in all_lessons:
        color = status_colors.get(lesson.status, '#6c757d')
        calendar_events.append({
            'title': f"{lesson.subject.name} - {lesson.teacher.user.last_name}",
            'start': f"{lesson.date}T{lesson.start_time}",
            'end': f"{lesson.date}T{lesson.end_time}",
            'url': f"/lesson/{lesson.id}/",
            'backgroundColor': color,
            'borderColor': color,
            'textColor': 'white'
        })

    # Групповые уроки
    for lesson in group_lessons:
        color = status_colors.get(lesson.status, '#6c757d')
        calendar_events.append({
            'title': f"👥 {lesson.subject.name} (группа)",
            'start': f"{lesson.date}T{lesson.start_time}",
            'end': f"{lesson.date}T{lesson.end_time}",
            'url': f"/student/group-lesson/{lesson.id}/",
            'backgroundColor': color,
            'borderColor': color,
            'textColor': 'white'
        })

    # ===== КОНТЕКСТ ДЛЯ ШАБЛОНА =====
    context = {
        # Основное
        'student': student,
        'user': user,

        # Финансы - используем get_balance()
        'balance': current_balance,  # ✅ Правильный баланс из метода
        'total_deposits': float(total_deposits),
        'total_expenses': float(total_expenses),
        'attended_cost': float(attended_cost_total),

        # Статистика по урокам
        'attended_lessons': attended_lessons,
        'debt_lessons': debt_lessons,
        'debt_cost': float(debt_cost_total),

        # Пополнения
        'recent_deposits': recent_deposits,

        # Уроки
        'upcoming_lessons': upcoming_lessons_list,
        'past_lessons': past_lessons,

        # Учителя
        'teachers': teachers,

        # Материалы
        'materials': materials,

        # Домашние задания
        'recent_homeworks': recent_homeworks,
        'homeworks': all_homeworks[:10],
        'stats': homework_stats,

        # Календарь
        'calendar_events': calendar_events,
        'now': timezone.now(),
    }

    # Отладка
    print(f"\n📊 КОНТЕКСТ ДЛЯ {user.username}:")
    print(f"   Баланс в БД: {float(user.balance)}")
    print(f"   Правильный баланс (get_balance): {current_balance}")
    print(f"   Проведено уроков на: {attended_cost_total}")
    print(f"   Домашних заданий: {homework_stats['total']}")

    return render(request, 'school/student/dashboard.html', context)

@login_required
def teacher_dashboard(request):
    """Личный кабинет учителя - с ДЗ"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    today = timezone.now().date()

    # ИСПОЛЬЗУЕМ TeacherFinanceHelper
    finance_helper = TeacherFinanceHelper(teacher)

    students = teacher.student_set.all().select_related('user')

    upcoming_lessons = Lesson.objects.filter(
        teacher=teacher,
        date__gte=today,
        status='scheduled'
    ).select_related('subject').order_by('date', 'start_time')[:20]

    today_lessons = Lesson.objects.filter(
        teacher=teacher,
        date=today,
        status='scheduled'
    ).select_related('subject').order_by('start_time')

    past_lessons = Lesson.objects.filter(
        teacher=teacher,
        status='completed'
    ).select_related('subject').order_by('-date')[:20]

    all_lessons = Lesson.objects.filter(
        teacher=teacher
    ).select_related(
        'subject'
    ).prefetch_related(
        Prefetch(
            'attendance',
            queryset=LessonAttendance.objects.select_related('student__user')
        )
    ).order_by('date', 'start_time')

    materials = Material.objects.filter(
        Q(teachers=teacher) | Q(created_by=request.user)
    ).distinct().order_by('-created_at')[:20]

    recent_payments = Payment.objects.filter(
        user=request.user,
        payment_type='teacher_payment'
    ).order_by('-created_at')[:10]

    # ✅ Добавляем ДЗ
    recent_homeworks = Homework.objects.filter(
        teacher=teacher
    ).select_related(
        'student__user', 'subject'
    ).prefetch_related('submission').order_by('-created_at')[:10]

    # Статистика по ДЗ
    all_homeworks = Homework.objects.filter(teacher=teacher)
    homework_stats = {
        'total': all_homeworks.count(),
        'pending': all_homeworks.filter(submission__isnull=True).count(),
        'submitted': all_homeworks.filter(submission__status='submitted').count(),
        'checked': all_homeworks.filter(submission__status='checked').count(),
        'overdue': all_homeworks.filter(
            deadline__lt=timezone.now(),
            submission__isnull=True
        ).count(),
    }

    # Календарь (существующий код)
    calendar_events = []
    for lesson in all_lessons:
        calc = LessonFinanceCalculator(lesson)
        stats = calc.stats

        # Определяем цвет
        if lesson.status == 'completed':
            bg_color = '#28a745'
        elif lesson.status == 'cancelled':
            bg_color = '#dc3545'
        elif lesson.status == 'overdue':
            bg_color = '#fd7e14'
        elif lesson.date < today and lesson.status == 'scheduled':
            bg_color = '#ffc107'
        elif lesson.date == today:
            bg_color = '#007bff'
        else:
            bg_color = '#6c757d'

        # Время начала
        time_str = lesson.start_time.strftime('%H:%M')

        # Получаем учеников
        students_att = lesson.attendance.all()
        total_count = students_att.count()

        if total_count == 0:
            title = ""
        else:
            names = []
            for attendance in students_att:
                student = attendance.student
                name = student.user.first_name or ""
                surname = student.user.last_name or ""

                if name and surname:
                    names.append(f"{name} {surname[0]}.")
                elif name:
                    names.append(name)
                elif surname:
                    names.append(surname)
                else:
                    names.append("Ученик")

            students_text = ", ".join(names)
            title = f"{time_str} {students_text}"

        calendar_events.append({
            'title': title,
            'start': f"{lesson.date}T{lesson.start_time}",
            'end': f"{lesson.date}T{lesson.end_time}",
            'url': f"/teacher/lesson/{lesson.id}/",
            'backgroundColor': bg_color,
            'borderColor': bg_color,
            'textColor': 'white',
            'finance': {
                'total_cost': stats['total_cost'],
                'teacher_payment': stats['teacher_payment']
            }
        })

    context = {
        'teacher': teacher,
        'finance': {
            'wallet_balance': float(finance_helper.wallet_balance),
            'payment_stats': finance_helper.get_payment_stats(30)
        },
        'students': students,
        'upcoming_lessons': upcoming_lessons,
        'today_lessons': today_lessons,
        'past_lessons': past_lessons,
        'all_lessons': all_lessons,
        'materials': materials,
        'recent_payments': recent_payments,
        'calendar_events': calendar_events,
        'recent_homeworks': recent_homeworks,
        'homework_stats': homework_stats,
        'all_homeworks': all_homeworks,

    }

    return render(request, 'school/teacher/dashboard.html', context)


@login_required
@require_GET
def api_student_completed_lessons(request):
    """API для получения завершенных уроков ученика по предмету"""
    if request.user.role != 'teacher':
        return JsonResponse({'error': 'Доступ запрещен'}, status=403)

    student_id = request.GET.get('student_id')
    subject_id = request.GET.get('subject_id')

    if not student_id or not subject_id:
        return JsonResponse({'error': 'Не указан student_id или subject_id'}, status=400)

    teacher = request.user.teacher_profile

    # Проверяем, что ученик принадлежит этому учителю
    student = get_object_or_404(Student, id=student_id, teachers=teacher)

    # Получаем завершенные уроки ученика по этому предмету
    completed_lessons = Lesson.objects.filter(
        teacher=teacher,
        subject_id=subject_id,
        attendance__student=student,
        status='completed'
    ).select_related('report').order_by('-date', '-start_time')

    lessons_data = []
    for lesson in completed_lessons:
        topic = lesson.report.topic if hasattr(lesson, 'report') and lesson.report else ''
        lessons_data.append({
            'id': lesson.id,
            'date': lesson.date.strftime('%d.%m.%Y'),
            'start_time': lesson.start_time.strftime('%H:%M'),
            'topic': topic,
        })

    return JsonResponse({'lessons': lessons_data})

@login_required
@require_POST
def teacher_homework_delete(request, homework_id):
    """Удаление домашнего задания"""
    if request.user.role != 'teacher':
        return JsonResponse({'error': 'Доступ запрещен'}, status=403)

    teacher = request.user.teacher_profile
    homework = get_object_or_404(Homework, id=homework_id, teacher=teacher)

    title = homework.title
    homework.delete()

    # ✅ ЛОГИРОВАНИЕ удаления
    log_user_action(
        request,
        'homework_delete',
        f'Удалено домашнее задание #{homework_id} - {title}',
        object_id=homework_id,
        object_type='homework'
    )

    return JsonResponse({'success': True, 'message': 'Задание удалено'})

@login_required
def teacher_lesson_detail(request, lesson_id):
    """Детальная страница урока для учителя"""
    lesson = get_object_or_404(Lesson, id=lesson_id)

    # ✅ СНАЧАЛА проверяем права доступа
    if request.user.role != 'teacher' or lesson.teacher.user != request.user:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    # ✅ ТОЛЬКО ПОСЛЕ ПРОВЕРКИ - логируем
    log_user_action(
        request,
        'lesson_view',
        f'Просмотр урока #{lesson.id} - {lesson.subject.name}',
        object_id=lesson.id,
        object_type='lesson',
        additional_data={
            'teacher': lesson.teacher.user.get_full_name(),
            'students_count': lesson.students.count(),
            'status': lesson.status
        }
    )

    # ИСПОЛЬЗУЕМ LessonFinanceCalculator
    calculator = LessonFinanceCalculator(lesson)
    stats = calculator.stats

    attendances = lesson.attendance.all().select_related('student__user')

    report = None
    if hasattr(lesson, 'report'):
        report = lesson.report

    form = None
    if lesson.status == 'scheduled':
        form = LessonReportForm()

    previous_lessons = Lesson.objects.filter(
        teacher=lesson.teacher,
        attendance__student__in=[a.student for a in attendances],
        date__lt=lesson.date,
        status='completed'
    ).distinct().order_by('-date')[:5]

    homeworks = Homework.objects.filter(lesson=lesson).order_by('-created_at')

    context = {
        'lesson': lesson,
        'attendances': calculator.get_attendance_details(),
        'finance': {
            'total_cost': stats['total_cost'],
            'teacher_payment': stats['teacher_payment'],
            'attended_cost': stats['attended_cost'],
            'attended_payment': stats['attended_payment'],
            'debt_cost': stats['debt_cost'],
            'students_total': stats['students_total'],
            'students_attended': stats['students_attended'],
            'students_debt': stats['students_debt']
        },
        'report': report,
        'form': form,
        'previous_lessons': previous_lessons,
        'homeworks': homeworks,
    }

    return render(request, 'school/teacher/lesson_detail.html', context)

@login_required
def lesson_detail(request, lesson_id):
    """Детальная страница урока для ученика"""

    lesson = get_object_or_404(Lesson, id=lesson_id)

    # ПРОВЕРКА НА ПРОСРОЧКУ
    from datetime import datetime
    if lesson.status == 'scheduled':
        lesson_datetime = datetime.combine(lesson.date, lesson.start_time)
        now = datetime.now()

        if lesson_datetime < now:
            lesson.status = 'overdue'
            lesson.save()

    user = request.user

    if user.role == 'student':
        try:
            attendance = lesson.attendance.get(student__user=user)
        except LessonAttendance.DoesNotExist:
            messages.error(request, 'Доступ запрещен')
            return redirect('dashboard')

        # ЛОГИРОВАНИЕ ПРОСМОТРА УРОКА
        log_user_action(
            request,
            'lesson_view',
            f'Просмотр урока #{lesson.id} - {lesson.subject.name}',
            object_id=lesson.id,
            object_type='lesson',
            additional_data={
                'teacher': lesson.teacher.user.get_full_name(),
                'subject': lesson.subject.name,
                'role': 'student'
            }
        )

        # ЛОГИРОВАНИЕ ВХОДА В ВИДЕО (если есть параметр)
        if lesson.video_room and request.GET.get('enter_video') == '1':
            log_user_action(
                request,
                'video_room_enter',
                f'Вход в видео-комнату урока #{lesson.id} - {lesson.subject.name}',
                object_id=lesson.id,
                object_type='lesson',
                additional_data={
                    'teacher': lesson.teacher.user.get_full_name(),
                    'subject': lesson.subject.name,
                    'role': 'student'
                }
            )

        # ПОЛУЧАЕМ ВСЕХ УЧАСТНИКОВ
        all_attendances = lesson.attendance.all().select_related('student__user')

    else:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    # ИСПОЛЬЗУЕМ LessonFinanceCalculator
    calculator = LessonFinanceCalculator(lesson)

    previous_lessons = []
    if user.role == 'student':
        previous_lessons = Lesson.objects.filter(
            teacher=lesson.teacher,
            attendance__student=attendance.student,
            date__lt=lesson.date
        ).distinct().order_by('-date', '-start_time')[:5]

    report = None
    if hasattr(lesson, 'report'):
        report = lesson.report

    # Обработка оценки урока
    if request.method == 'POST' and user.role == 'student' and lesson.status == 'completed' and not hasattr(lesson, 'feedback'):
        rating = request.POST.get('rating')
        comment = request.POST.get('comment', '')
        is_public = request.POST.get('is_public') == 'on'

        if rating and rating.isdigit():
            feedback = LessonFeedback.objects.create(
                lesson=lesson,
                student=attendance.student,
                teacher=lesson.teacher,
                rating=int(rating),
                comment=comment,
                is_public=is_public
            )

            teacher_rating, created = TeacherRating.objects.get_or_create(teacher=lesson.teacher)
            teacher_rating.update_stats()

            messages.success(request, 'Спасибо за вашу оценку!')
            return redirect('lesson_detail', lesson_id=lesson.id)
        else:
            messages.error(request, 'Пожалуйста, поставьте оценку')

    context = {
        'lesson': lesson,
        'attendance': attendance,
        'attendances': all_attendances,
        'attendance_details': calculator.get_attendance_details(),
        'finance': {
            'student_cost': float(attendance.cost),
            'total_cost': calculator.stats['total_cost'],
            'students_total': calculator.stats['students_total']
        },
        'report': report,
        'previous_lessons': previous_lessons,
    }

    return render(request, 'school/student/lesson_detail.html', context)

@login_required
@require_POST
def log_video_entry(request, lesson_id):
    """Логирование входа в видео-комнату"""
    try:
        lesson = get_object_or_404(Lesson, id=lesson_id)

        log_user_action(
            request,
            'video_room_enter',
            f'Вход в видео-комнату урока #{lesson.id} - {lesson.subject.name}',
            object_id=lesson.id,
            object_type='lesson',
            additional_data={
                'teacher': lesson.teacher.user.get_full_name(),
                'subject': lesson.subject.name,
                'role': 'student',
                'room': lesson.video_room
            }
        )

        return JsonResponse({'status': 'ok'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@staff_member_required
@require_POST
def admin_complete_lesson(request, lesson_id):
    """Завершает занятие из админ-панели - ПОЛНЫЙ РЕФАКТОРИНГ"""
    print(f"REQUEST METHOD: {request.method}")
    print(f"REQUEST POST: {request.POST}")
    print(f"LESSON ID: {lesson_id}")
    print(f"\n{'🔥' * 30}")
    print(f"🔥🔥🔥 ЗАВЕРШЕНИЕ УРОКА #{lesson_id} 🔥🔥🔥")
    print(f"{'🔥' * 30}\n")

    try:
        lesson = Lesson.objects.select_related('teacher__user', 'subject').get(pk=lesson_id)

        if lesson.status == 'completed':
            messages.error(request, 'Занятие уже завершено')
            return redirect('admin:school_lesson_change', lesson_id)

        # ИСПОЛЬЗУЕМ LessonFinanceCalculator
        calculator = LessonFinanceCalculator(lesson)
        stats = calculator.stats

        if stats['students_total'] == 0:
            messages.error(request, 'Нет учеников на уроке')
            return redirect('admin:school_lesson_change', lesson_id)

        # Проверяем POST данные
        report_data = {
            'topic': request.POST.get('topic', '').strip(),
            'covered_material': request.POST.get('covered_material', '').strip(),
            'homework': request.POST.get('homework', '').strip(),
            'student_progress': request.POST.get('student_progress', '').strip(),
            'next_lesson_plan': request.POST.get('next_lesson_plan', '').strip()
        }

        required_fields = ['topic', 'covered_material', 'homework', 'student_progress']
        missing = [f for f in required_fields if not report_data[f]]
        if missing:
            messages.error(request, f'Заполните обязательные поля: {", ".join(missing)}')
            return redirect('admin:school_lesson_change', lesson_id)

        with transaction.atomic():
            processed_students = []

            for attendance in calculator.attendances:
                student = attendance.student
                user = student.user

                # ✅ ЗАПОМИНАЕМ БАЛАНС ДО СПИСАНИЯ
                old_balance = float(user.balance)

                # ✅ СПИСЫВАЕМ ДЕНЬГИ С БАЛАНСА УЧЕНИКА
                user.balance -= attendance.cost
                user.save()

                # УРОК СЧИТАЕТСЯ ПРОВЕДЕННЫМ
                attendance.status = 'attended'
                attendance.save()

                # СОЗДАЕМ ЗАПИСЬ О ПЛАТЕЖЕ
                Payment.objects.create(
                    user=user,
                    amount=attendance.cost,
                    payment_type='expense',
                    description=f'Оплата занятия {lesson.date} ({lesson.subject.name})',
                    lesson=lesson
                )

                student_data = {
                    'name': user.get_full_name(),
                    'cost': float(attendance.cost),
                    'teacher_payment': float(attendance.teacher_payment_share),
                    'old_balance': old_balance,
                    'new_balance': float(user.balance),
                    'debt': False
                }
                processed_students.append(student_data)

                print(f"💰 Баланс ученика {user.username}: {old_balance} → {user.balance} (списано {attendance.cost})")

            # НАЧИСЛЯЕМ УЧИТЕЛЮ
            old_teacher_balance = lesson.teacher.wallet_balance
            lesson.teacher.wallet_balance += calculator.teacher_payment
            lesson.teacher.save()

            if calculator.teacher_payment > 0:
                Payment.objects.create(
                    user=lesson.teacher.user,
                    amount=calculator.teacher_payment,
                    payment_type='teacher_payment',
                    description=f'Выплата за урок {lesson.date} ({lesson.subject.name})',
                    lesson=lesson
                )

            # МЕНЯЕМ СТАТУС УРОКА
            lesson.status = 'completed'
            lesson.save()

            # ✅ СОЗДАЕМ ИЛИ ОБНОВЛЯЕМ ОТЧЕТ (ВНУТРИ ТРАНЗАКЦИИ)
            report, created = LessonReport.objects.update_or_create(
                lesson=lesson,
                defaults={
                    'topic': report_data['topic'],
                    'covered_material': report_data['covered_material'],
                    'homework': report_data['homework'],
                    'student_progress': report_data['student_progress'],
                    'next_lesson_plan': report_data['next_lesson_plan']
                }
            )

            if created:
                print(f"✅ Создан новый отчет #{report.id}")
            else:
                print(f"✅ Обновлен существующий отчет #{report.id}")

        # ✅ ОТПРАВЛЯЕМ УВЕДОМЛЕНИЕ В TELEGRAM
        try:
            from school.telegram import notify_lesson_completed
            notify_lesson_completed(lesson, report)
        except Exception as e:
            print(f"❌ Ошибка отправки Telegram уведомления: {e}")

        messages.success(request, f'✅ Урок успешно завершен! Отчет #{report.id} создан.')
        return redirect('admin:school_lesson_change', lesson_id)

    except Lesson.DoesNotExist:
        messages.error(request, 'Занятие не найдено')
        return redirect('admin:school_lesson_changelist')
    except Exception as e:
        print(f"❌ ОШИБКА: {e}")
        traceback.print_exc()
        messages.error(request, f'Ошибка: {str(e)}')
        return redirect('admin:school_lesson_change', lesson_id)


@staff_member_required
def admin_finance_dashboard(request):
    """Финансовый дашборд для администратора - ПОЛНЫЙ РЕФАКТОРИНГ"""

    today = timezone.now().date()
    start_date = request.GET.get('start_date', today.replace(day=1).strftime('%Y-%m-%d'))
    end_date = request.GET.get('end_date', today.strftime('%Y-%m-%d'))

    start = datetime.strptime(start_date, '%Y-%m-%d').date()
    end = datetime.strptime(end_date, '%Y-%m-%d').date()

    # Получаем данные за период
    lessons = Lesson.objects.filter(date__gte=start, date__lte=end)
    payments = Payment.objects.filter(created_at__date__gte=start, created_at__date__lte=end)

    # ИСПОЛЬЗУЕМ PeriodFinanceCalculator
    period_calc = PeriodFinanceCalculator(lessons, payments)

    # Статистика по ученикам
    students_with_debt = Student.objects.filter(user__balance__lt=0).count()
    total_debt = abs(
        Student.objects.filter(user__balance__lt=0).aggregate(Sum('user__balance'))['user__balance__sum'] or 0)

    students_with_balance = Student.objects.filter(user__balance__gt=0).count()
    total_balance = Student.objects.filter(user__balance__gt=0).aggregate(Sum('user__balance'))[
                        'user__balance__sum'] or 0

    # Топ-10 учеников
    top_students = Student.objects.select_related('user').order_by('-user__balance')[:10]
    top_debtors = Student.objects.filter(user__balance__lt=0).select_related('user').order_by('user__balance')[:10]

    # Статистика по учителям
    teachers_total_balance = Teacher.objects.aggregate(Sum('wallet_balance'))['wallet_balance__sum'] or 0

    context = {
        'period': {
            'start': start_date,
            'end': end_date,
            'start_formatted': start.strftime('%d.%m.%Y'),
            'end_formatted': end.strftime('%d.%m.%Y'),
        },
        'lessons_stats': period_calc.lessons_stats,
        'payments_stats': period_calc.payments_stats,
        'school_finances': period_calc.school_finances,
        'daily_stats': period_calc.get_daily_stats(start, end),
        'students': {
            'with_debt': students_with_debt,
            'total_debt': float(total_debt),
            'with_balance': students_with_balance,
            'total_balance': float(total_balance),
            'top_students': [
                {
                    'name': s.user.get_full_name(),
                    'balance': float(s.user.balance)
                } for s in top_students
            ],
            'top_debtors': [
                {
                    'name': s.user.get_full_name(),
                    'debt': float(s.user.balance)
                } for s in top_debtors
            ]
        },
        'teachers': {
            'total_balance': float(teachers_total_balance)
        }
    }

    return render(request, 'admin/finance/dashboard.html', context)


# Остальные views с аналогичным рефакторингом...
# (здесь идут все остальные функции, но я их пропускаю для краткости,
# так как принцип везде одинаковый - использовать созданные helper-классы)

@login_required
def student_deposit(request):
    """Пополнение баланса ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    if request.method == 'POST':
        amount = request.POST.get('amount')
        description = request.POST.get('description', 'Пополнение счета')

        try:
            amount = Decimal(amount)
            if amount <= 0:
                messages.error(request, 'Сумма должна быть положительной')
                return redirect('student_dashboard')

            student = request.user.student_profile

            deposit = Deposit.objects.create(
                student=student,
                amount=amount,
                description=description,
                created_by=request.user
            )

            request.user.balance += amount
            request.user.save()

            messages.success(request, f'Счет пополнен на {amount} ₽')

        except (ValueError, TypeError, Decimal.InvalidOperation):
            messages.error(request, 'Неверная сумма')

        return redirect('student_dashboard')

    return redirect('student_dashboard')


# ============================================
# ЧАСТЬ 4: ФУНКЦИИ ЭКСПОРТА
# ============================================

@staff_member_required
def export_teacher_payment(request, format, teacher_id, start_date, end_date):
    """Экспорт отчета в разных форматах"""
    teacher = get_object_or_404(Teacher, id=teacher_id)
    start = datetime.strptime(start_date, '%Y-%m-%d').date()
    end = datetime.strptime(end_date, '%Y-%m-%d').date()

    lessons = Lesson.objects.filter(
        teacher=teacher,
        status='completed',
        date__gte=start,
        date__lte=end
    ).prefetch_related('attendance__student__user', 'subject').order_by('date')

    # ИСПОЛЬЗУЕМ PeriodFinanceCalculator
    period_calc = PeriodFinanceCalculator(lessons)
    stats = period_calc.lessons_stats

    if format == 'excel':
        return export_to_excel(teacher, lessons, start, end, stats['teacher_payment'])
    elif format == 'word':
        return export_to_word(teacher, lessons, start, end, stats['teacher_payment'])
    elif format == 'pdf':
        return export_to_pdf(teacher, lessons, start, end, stats['teacher_payment'])
    else:
        return HttpResponse('Неподдерживаемый формат', status=400)


def export_to_excel(teacher, lessons, start, end, total_payment):
    """Экспорт в Excel"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Расчет выплат"

    title_font = Font(name='Arial', size=14, bold=True)
    header_font = Font(name='Arial', size=11, bold=True)
    normal_font = Font(name='Arial', size=10)
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font_white = Font(name='Arial', size=11, bold=True, color="FFFFFF")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))

    ws.merge_cells('A1:F1')
    cell = ws['A1']
    cell.value = f"Расчет выплат учителю: {teacher.user.get_full_name()}"
    cell.font = title_font
    cell.alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:F2')
    cell = ws['A2']
    cell.value = f"Период: {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}"
    cell.font = normal_font
    cell.alignment = Alignment(horizontal='center')

    ws.append([])

    headers = ['Дата', 'Ученик', 'Предмет', 'Стоимость урока', 'Выплата учителю', 'Статус']
    ws.append(headers)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border

    row = 5
    for lesson in lessons:
        # ИСПОЛЬЗУЕМ LessonFinanceCalculator для детализации
        calculator = LessonFinanceCalculator(lesson)
        for attendance in calculator.get_attendance_details():
            ws.cell(row=row, column=1, value=lesson.date.strftime('%d.%m.%Y')).border = thin_border
            ws.cell(row=row, column=2, value=attendance['student_name']).border = thin_border
            ws.cell(row=row, column=3, value=lesson.subject.name).border = thin_border
            ws.cell(row=row, column=4, value=attendance['cost']).border = thin_border
            ws.cell(row=row, column=5, value=attendance['teacher_payment']).border = thin_border
            ws.cell(row=row, column=6, value=lesson.get_status_display()).border = thin_border

            ws.cell(row=row, column=4).number_format = '#,##0.00 ₽'
            ws.cell(row=row, column=5).number_format = '#,##0.00 ₽'
            row += 1

    row += 1
    ws.cell(row=row, column=4, value="ИТОГО:").font = header_font
    ws.cell(row=row, column=5, value=float(total_payment)).font = header_font
    ws.cell(row=row, column=5).number_format = '#,##0.00 ₽'

    column_widths = [12, 30, 20, 15, 15, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f"teacher_payment_{teacher.id}_{start}_{end}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    wb.save(response)
    return response


def export_to_word(teacher, lessons, start, end, total_payment):
    """Экспорт в Word"""
    doc = Document()

    title = doc.add_heading('Расчет выплат учителю', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Информация об учителе:', level=1)
    doc.add_paragraph(f'ФИО: {teacher.user.get_full_name()}')
    doc.add_paragraph(f'Email: {teacher.user.email}')
    doc.add_paragraph(f'Телефон: {teacher.user.phone}')

    doc.add_heading('Период расчета:', level=1)
    doc.add_paragraph(f'с {start.strftime("%d.%m.%Y")} по {end.strftime("%d.%m.%Y")}')

    doc.add_heading('Детализация уроков:', level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    headers = ['Дата', 'Ученик', 'Предмет', 'Стоимость', 'Выплата']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True

    for lesson in lessons:
        calculator = LessonFinanceCalculator(lesson)
        for attendance in calculator.get_attendance_details():
            row_cells = table.add_row().cells
            row_cells[0].text = lesson.date.strftime('%d.%m.%Y')
            row_cells[1].text = attendance['student_name']
            row_cells[2].text = lesson.subject.name
            row_cells[3].text = f"{attendance['cost']:.2f} ₽"
            row_cells[4].text = f"{attendance['teacher_payment']:.2f} ₽"

    doc.add_paragraph()
    total_para = doc.add_paragraph()
    total_para.add_run('ИТОГО К ВЫПЛАТЕ: ').bold = True
    total_para.add_run(f'{total_payment:.2f} ₽').bold = True

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="teacher_payment_{teacher.id}_{start}_{end}.docx"'

    doc.save(response)
    return response


def export_to_pdf(teacher, lessons, start, end, total_payment):
    """Экспорт в PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()

    title = Paragraph(f"Расчет выплат учителю", styles['Title'])
    elements.append(title)
    elements.append(Paragraph(f"<b>{teacher.user.get_full_name()}</b>", styles['Normal']))
    elements.append(Paragraph(f"Период: {start.strftime('%d.%m.%Y')} - {end.strftime('%d.%m.%Y')}", styles['Normal']))
    elements.append(Paragraph("<br/>", styles['Normal']))

    data = [['Дата', 'Ученик', 'Предмет', 'Стоимость', 'Выплата']]

    for lesson in lessons:
        calculator = LessonFinanceCalculator(lesson)
        for attendance in calculator.get_attendance_details():
            data.append([
                lesson.date.strftime('%d.%m.%Y'),
                attendance['student_name'],
                lesson.subject.name,
                f"{attendance['cost']:.2f} ₽",
                f"{attendance['teacher_payment']:.2f} ₽"
            ])

    data.append(['', '', '', 'ИТОГО:', f"{total_payment:.2f} ₽"])

    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -2), 1, colors.black),
        ('GRID', (0, -1), (-1, -1), 1, colors.black),
    ]))

    elements.append(table)
    doc.build(elements)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="teacher_payment_{teacher.id}_{start}_{end}.pdf"'

    return response


@staff_member_required
def admin_lesson_export(request, format):
    """Экспорт уроков из админки"""
    lessons = Lesson.objects.all().select_related(
        'teacher__user', 'subject', 'format'
    ).prefetch_related('attendance__student__user').order_by('-date', 'start_time')

    teacher_id = request.GET.get('teacher__id__exact')
    student_id = request.GET.get('student__id__exact')
    subject_id = request.GET.get('subject__id__exact')
    status = request.GET.get('status__exact')
    date_from = request.GET.get('date__gte')
    date_to = request.GET.get('date__lte')

    if teacher_id:
        lessons = lessons.filter(teacher_id=teacher_id)
    if student_id:
        lessons = lessons.filter(attendance__student_id=student_id)
    if subject_id:
        lessons = lessons.filter(subject_id=subject_id)
    if status:
        lessons = lessons.filter(status=status)
    if date_from:
        lessons = lessons.filter(date__gte=date_from)
    if date_to:
        lessons = lessons.filter(date__lte=date_to)

    # ИСПОЛЬЗУЕМ PeriodFinanceCalculator
    period_calc = PeriodFinanceCalculator(lessons)
    stats = period_calc.lessons_stats

    title = f"Экспорт уроков"

    if format == 'excel':
        return export_lessons_excel(lessons, title, stats['completed'], stats['cancelled'],
                                    stats['overdue'], stats['total_cost'], stats['teacher_payment'])
    elif format == 'csv':
        return export_lessons_csv(lessons, title, stats['completed'], stats['cancelled'],
                                  stats['overdue'], stats['total_cost'], stats['teacher_payment'])
    elif format == 'html':
        return export_lessons_html(lessons, title, stats['completed'], stats['cancelled'],
                                   stats['overdue'], stats['total_cost'], stats['teacher_payment'])
    elif format == 'pdf':
        return export_lessons_pdf(lessons, title, stats['completed'], stats['cancelled'],
                                  stats['overdue'], stats['total_cost'], stats['teacher_payment'])
    else:
        messages.error(request, 'Неподдерживаемый формат')
        return redirect(request.META.get('HTTP_REFERER', 'admin:school_lesson_changelist'))


def export_lessons_excel(lessons, title, completed_count, cancelled_count, overdue_count, total_cost, total_payment):
    """Экспорт уроков в Excel"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Уроки"

    title_font = Font(name='Arial', size=16, bold=True)
    header_font = Font(name='Arial', size=12, bold=True)
    header_fill = PatternFill(start_color="417690", end_color="417690", fill_type="solid")
    header_font_white = Font(name='Arial', size=12, bold=True, color="FFFFFF")
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))

    ws.merge_cells('A1:I1')
    cell = ws['A1']
    cell.value = title
    cell.font = title_font
    cell.alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:I2')
    cell = ws['A2']
    cell.value = f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    cell.font = Font(italic=True)
    cell.alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:I3')
    cell = ws['A3']
    cell.value = f"Всего: {lessons.count()} | Проведено: {completed_count} | Отменено: {cancelled_count} | Просрочено: {overdue_count} | Сумма: {total_cost:,.2f} ₽ | Выплаты: {total_payment:,.2f} ₽"
    cell.font = Font(bold=True)
    cell.alignment = Alignment(horizontal='center')

    ws.append([])

    headers = ['ID урока', 'Дата', 'Время', 'Учитель', 'Ученик', 'Предмет', 'Стоимость', 'Выплата учителю', 'Статус']

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col)
        cell.value = header
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border

    row = 6
    for lesson in lessons:
        calculator = LessonFinanceCalculator(lesson)
        for attendance in calculator.get_attendance_details():
            ws.cell(row=row, column=1, value=lesson.id).border = thin_border
            ws.cell(row=row, column=2, value=lesson.date.strftime('%d.%m.%Y')).border = thin_border
            ws.cell(row=row, column=3,
                    value=f"{lesson.start_time.strftime('%H:%M')}-{lesson.end_time.strftime('%H:%M')}").border = thin_border
            ws.cell(row=row, column=4, value=lesson.teacher.user.get_full_name()).border = thin_border
            ws.cell(row=row, column=5, value=attendance['student_name']).border = thin_border
            ws.cell(row=row, column=6, value=lesson.subject.name).border = thin_border
            ws.cell(row=row, column=7, value=attendance['cost']).border = thin_border
            ws.cell(row=row, column=8, value=attendance['teacher_payment']).border = thin_border
            ws.cell(row=row, column=9, value=lesson.get_status_display()).border = thin_border

            ws.cell(row=row, column=7).number_format = '#,##0.00 ₽'
            ws.cell(row=row, column=8).number_format = '#,##0.00 ₽'

            status_cell = ws.cell(row=row, column=9)
            if lesson.status == 'completed':
                status_cell.font = Font(color="28A745", bold=True)
            elif lesson.status == 'cancelled':
                status_cell.font = Font(color="DC3545", bold=True)
            elif lesson.status == 'overdue':
                status_cell.font = Font(color="FFC107", bold=True)
            elif lesson.status == 'scheduled':
                status_cell.font = Font(color="007BFF", bold=True)

            row += 1

    row += 1
    ws.cell(row=row, column=6, value="ИТОГО:").font = Font(bold=True)
    ws.cell(row=row, column=7, value=float(total_cost)).font = Font(bold=True)
    ws.cell(row=row, column=7).number_format = '#,##0.00 ₽'
    ws.cell(row=row, column=8, value=float(total_payment)).font = Font(bold=True)
    ws.cell(row=row, column=8).number_format = '#,##0.00 ₽'

    column_widths = [8, 12, 15, 25, 25, 20, 15, 18, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f"lessons_export_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    wb.save(response)
    return response


def export_lessons_csv(lessons, title, completed_count, cancelled_count, overdue_count, total_cost, total_payment):
    """Экспорт уроков в CSV"""
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    filename = f"lessons_export_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    response.write('\ufeff')
    writer = csv.writer(response, delimiter=';')

    writer.writerow([title])
    writer.writerow([f"Дата экспорта: {datetime.now().strftime('%d.%m.%Y %H:%M')}"])
    writer.writerow([
        f"Всего: {lessons.count()} | Проведено: {completed_count} | Отменено: {cancelled_count} | Просрочено: {overdue_count}"])
    writer.writerow([f"Общая стоимость: {total_cost:.2f} ₽ | Общая сумма выплат: {total_payment:.2f} ₽"])
    writer.writerow([])

    writer.writerow(['ID', 'Дата', 'Время', 'Учитель', 'Ученик', 'Предмет', 'Стоимость', 'Выплата учителю', 'Статус'])

    for lesson in lessons:
        calculator = LessonFinanceCalculator(lesson)
        for attendance in calculator.get_attendance_details():
            writer.writerow([
                lesson.id,
                lesson.date.strftime('%d.%m.%Y'),
                f"{lesson.start_time.strftime('%H:%M')}-{lesson.end_time.strftime('%H:%M')}",
                lesson.teacher.user.get_full_name(),
                attendance['student_name'],
                lesson.subject.name,
                f"{attendance['cost']:.2f}",
                f"{attendance['teacher_payment']:.2f}",
                lesson.get_status_display(),
            ])

    return response


def export_lessons_html(lessons, title, completed_count, cancelled_count, overdue_count, total_cost, total_payment):
    """Экспорт уроков в HTML"""
    context = {
        'title': title,
        'lessons': lessons,
        'export_date': datetime.now().strftime('%d.%m.%Y %H:%M'),
        'completed_count': completed_count,
        'cancelled_count': cancelled_count,
        'overdue_count': overdue_count,
        'total_cost': total_cost,
        'total_payment': total_payment,
    }

    html_content = render_to_string('admin/school/lesson/export.html', context)

    response = HttpResponse(html_content, content_type='text/html; charset=utf-8')
    filename = f"lessons_export_{datetime.now().strftime('%Y%m%d_%H%M')}.html"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


def export_lessons_pdf(lessons, title, completed_count, cancelled_count, overdue_count, total_cost, total_payment):
    """Экспорт уроков в PDF"""
    context = {
        'title': title,
        'lessons': lessons,
        'export_date': datetime.now().strftime('%d.%m.%Y %H:%M'),
        'completed_count': completed_count,
        'cancelled_count': cancelled_count,
        'overdue_count': overdue_count,
        'total_cost': total_cost,
        'total_payment': total_payment,
        'pdf_mode': True,
    }

    html_string = render_to_string('admin/school/lesson/export.html', context)

    response = HttpResponse(content_type='application/pdf')
    filename = f"lessons_export_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    HTML(string=html_string).write_pdf(response)

    return response


def download_import_template(request):
    """Скачать шаблон для импорта с поддержкой ID"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="import_lessons_template.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Импорт уроков"

    headers = [
        'Дата', 'Время начала', 'Время окончания',
        'ID учителя', 'Учитель (ФИО)',
        'ID учеников', 'Ученики (ФИО через ;)',
        'Предмет', 'Стоимость урока', 'Выплата учителю', 'Статус'
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill(start_color="417690", end_color="417690", fill_type="solid")

    examples = [
        ['01.03.2026', '10:00', '11:00', '10', 'Иванов Иван', '13', 'Петров Петр', 'Математика', '1000', '500',
         'scheduled'],
        ['02.03.2026', '11:00', '12:00', '11', 'Петрова Анна', '14;15', 'Сидоров Сидор; Козлова Елена', 'Русский язык',
         '1500', '900', 'scheduled'],
        ['03.03.2026', '14:00', '15:00', '12', 'Смирнов Павел', '16;17;18',
         'Соколов Максим; Волкова Дарья; Морозов Алексей', 'Английский язык', '2400', '1500', 'scheduled'],
    ]

    for row_num, example in enumerate(examples, start=2):
        for col_num, value in enumerate(example, 1):
            ws.cell(row=row_num, column=col_num, value=value)

    column_widths = [12, 15, 15, 12, 25, 15, 30, 20, 15, 15, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    wb.save(response)
    return response

@login_required
def export_calendar_pdf(request):
    """Экспорт календаря в PDF с фильтром по дате"""
    user = request.user
    
    # Получаем даты из GET-параметров
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    view_type = request.GET.get('view', 'month')  # month, week, day
    
    # Если даты не указаны, берем текущий месяц
    if not start_date_str or not end_date_str:
        today = timezone.now().date()
        start_date = today.replace(day=1)
        if today.month == 12:
            end_date = today.replace(year=today.year+1, month=1, day=1) - timedelta(days=1)
        else:
            end_date = today.replace(month=today.month+1, day=1) - timedelta(days=1)
    else:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    
    # Получаем уроки за период
    if user.role == 'student':
        student = user.student_profile
        lessons = Lesson.objects.filter(
            attendance__student=student,
            date__gte=start_date,
            date__lte=end_date
        ).select_related('teacher__user', 'subject').order_by('date', 'start_time')
    else:
        lessons = Lesson.objects.none()
    
    # Группируем по датам
    lessons_by_date = {}
    for lesson in lessons:
        date_str = lesson.date.strftime('%Y-%m-%d')
        if date_str not in lessons_by_date:
            lessons_by_date[date_str] = []
        
        # Формат: "10:00 Английский - Ива"
        time_str = lesson.start_time.strftime('%H:%M')
        teacher_short = lesson.teacher.user.last_name
        
        lessons_by_date[date_str].append({
            'time': time_str,
            'subject': lesson.subject.name,
            'teacher': teacher_short,
            'full': f"{time_str} {lesson.subject.name} - {teacher_short}"
        })
    
    # Создаем сетку календаря
    calendar_days = []
    current = start_date
    week = []
    
    # Добавляем пустые ячейки для первого дня
    first_weekday = start_date.weekday()
    for _ in range(first_weekday):
        week.append({'day': '', 'lessons': [], 'is_current_month': False})
    
    # Заполняем все дни
    while current <= end_date:
        if len(week) == 7:
            calendar_days.append(week)
            week = []
        
        date_str = current.strftime('%Y-%m-%d')
        week.append({
            'day': current.day,
            'lessons': lessons_by_date.get(date_str, []),
            'is_current_month': True,
            'date': current
        })
        current += timedelta(days=1)
    
    # Добавляем пустые ячейки в конец
    while len(week) < 7:
        week.append({'day': '', 'lessons': [], 'is_current_month': False})
    if week:
        calendar_days.append(week)
    
    context = {
        'user': user,
        'calendar_days': calendar_days,
        'start_date': start_date.strftime('%d.%m.%Y'),
        'end_date': end_date.strftime('%d.%m.%Y'),
        'month_name': ['январь', 'февраль', 'март', 'апрель', 'май', 'июнь',
                      'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь'][start_date.month-1],
        'year': start_date.year,
        'export_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
        'lessons_count': sum(len(v) for v in lessons_by_date.values()),
    }
    
    # Рендерим PDF
    html_string = render_to_string('school/student/calendar_month_pdf.html', context)
    
    response = HttpResponse(content_type='application/pdf')
    filename = f"calendar_{start_date.strftime('%Y%m')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(response)
    return response


@login_required
def teacher_export_calendar_pdf(request):
    """Экспорт календаря учителя в PDF с выбором месяца"""
    log_user_action(
        request,
        'calendar_export',
        f'Экспорт календаря: {request.GET.get("month")}.{request.GET.get("year")}',
        additional_data={
            'month': request.GET.get('month'),
            'year': request.GET.get('year'),
            'view': request.GET.get('view', 'month')
        }
    )
    user = request.user

    if user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = user.teacher_profile

    # Получаем месяц и год из GET-параметров
    month = request.GET.get('month')
    year = request.GET.get('year')

    # Если не указаны, берем текущий месяц
    today = timezone.now().date()
    if not month or not year:
        month = today.month
        year = today.year
    else:
        month = int(month)
        year = int(year)

    # Определяем первый и последний день месяца
    start_date = date(year, month, 1)
    if month == 12:
        end_date = date(year + 1, 1, 1) - timedelta(days=1)
    else:
        end_date = date(year, month + 1, 1) - timedelta(days=1)

    # Получаем все уроки учителя за месяц
    lessons = Lesson.objects.filter(
        teacher=teacher,
        date__gte=start_date,
        date__lte=end_date
    ).select_related('subject').prefetch_related('attendance__student__user').order_by('date', 'start_time')

    # Получаем групповые уроки
    group_lessons = GroupLesson.objects.filter(
        teacher=teacher,
        date__gte=start_date,
        date__lte=end_date
    ).select_related('subject').prefetch_related('enrollments__student__user')

    # Группируем по датам
    lessons_by_date = {}

    # Обычные уроки
    for lesson in lessons:
        date_str = lesson.date.strftime('%Y-%m-%d')
        if date_str not in lessons_by_date:
            lessons_by_date[date_str] = []

        # Формируем список учеников
        students_list = [a.student.user.get_full_name() for a in lesson.attendance.all()]

        time_str = lesson.start_time.strftime('%H:%M')

        lessons_by_date[date_str].append({
            'time': time_str,
            'subject': lesson.subject.name,
            'students': students_list,
            'students_count': len(students_list),
            'full': f"{time_str} {lesson.subject.name} ({len(students_list)} уч.)",
            'type': 'individual'
        })

    # Групповые уроки
    for lesson in group_lessons:
        date_str = lesson.date.strftime('%Y-%m-%d')
        if date_str not in lessons_by_date:
            lessons_by_date[date_str] = []

        students_list = [e.student.user.get_full_name() for e in lesson.enrollments.all()]
        time_str = lesson.start_time.strftime('%H:%M')

        lessons_by_date[date_str].append({
            'time': time_str,
            'subject': lesson.subject.name,
            'students': students_list,
            'students_count': len(students_list),
            'full': f"👥 {time_str} {lesson.subject.name} ({len(students_list)} уч.)",
            'type': 'group'
        })

    # Создаем сетку календаря
    calendar_days = []
    current = start_date
    week = []

    # Добавляем пустые ячейки для первого дня
    first_weekday = start_date.weekday()
    for _ in range(first_weekday):
        week.append({'day': '', 'lessons': [], 'is_current_month': False})

    # Заполняем все дни
    while current <= end_date:
        if len(week) == 7:
            calendar_days.append(week)
            week = []

        date_str = current.strftime('%Y-%m-%d')
        week.append({
            'day': current.day,
            'lessons': lessons_by_date.get(date_str, []),
            'is_current_month': True,
            'date': current
        })
        current += timedelta(days=1)

    # Добавляем пустые ячейки в конец
    while len(week) < 7:
        week.append({'day': '', 'lessons': [], 'is_current_month': False})
    if week:
        calendar_days.append(week)

    # Список месяцев для выпадающего списка
    months = [
        {'value': 1, 'name': 'Январь'},
        {'value': 2, 'name': 'Февраль'},
        {'value': 3, 'name': 'Март'},
        {'value': 4, 'name': 'Апрель'},
        {'value': 5, 'name': 'Май'},
        {'value': 6, 'name': 'Июнь'},
        {'value': 7, 'name': 'Июль'},
        {'value': 8, 'name': 'Август'},
        {'value': 9, 'name': 'Сентябрь'},
        {'value': 10, 'name': 'Октябрь'},
        {'value': 11, 'name': 'Ноябрь'},
        {'value': 12, 'name': 'Декабрь'},
    ]

    # Список годов (текущий год и +/- 2 года)
    current_year = timezone.now().year
    years = [current_year - 2, current_year - 1, current_year, current_year + 1, current_year + 2]

    context = {
        'user': user,
        'teacher': teacher,
        'calendar_days': calendar_days,
        'month_name': months[month - 1]['name'],
        'month': month,
        'year': year,
        'months': months,
        'years': years,
        'export_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
        'lessons_count': sum(len(v) for v in lessons_by_date.values()),
    }

    # Рендерим PDF
    html_string = render_to_string('school/teacher/calendar_month_pdf.html', context)

    response = HttpResponse(content_type='application/pdf')
    filename = f"teacher_calendar_{year}_{month:02d}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(response)
    return response


@staff_member_required
def admin_export_calendar_pdf(request):
    """Экспорт календаря админки в PDF с выбором месяца"""

    # Получаем месяц и год из GET-параметров
    month = request.GET.get('month')
    year = request.GET.get('year')

    today = timezone.now().date()

    # Если не указаны, берем текущий месяц
    if not month or not year:
        month = today.month
        year = today.year
    else:
        month = int(month)
        year = int(year)

    # Определяем первый и последний день месяца
    start_date = date(year, month, 1)
    if month == 12:
        end_date = date(year + 1, 1, 1) - timedelta(days=1)
    else:
        end_date = date(year, month + 1, 1) - timedelta(days=1)

    # Получаем все уроки за месяц
    lessons = Lesson.objects.filter(
        date__gte=start_date,
        date__lte=end_date
    ).select_related(
        'teacher__user', 'subject'
    ).prefetch_related(
        'attendance__student__user'
    ).order_by('date', 'start_time')

    # Групповые уроки
    group_lessons = GroupLesson.objects.filter(
        date__gte=start_date,
        date__lte=end_date
    ).select_related(
        'teacher__user', 'subject'
    ).prefetch_related(
        'enrollments__student__user'
    )

    # Группируем по датам
    lessons_by_date = {}

    # Обычные уроки
    for lesson in lessons:
        date_str = lesson.date.strftime('%Y-%m-%d')
        if date_str not in lessons_by_date:
            lessons_by_date[date_str] = []

        # Список учеников
        students_list = [a.student.user.get_full_name() for a in lesson.attendance.all()]
        time_str = lesson.start_time.strftime('%H:%M')

        lessons_by_date[date_str].append({
            'time': time_str,
            'subject': lesson.subject.name,
            'teacher': lesson.teacher.user.last_name,
            'students_count': len(students_list),
            'full': f"{time_str} {lesson.subject.name} - {lesson.teacher.user.last_name} ({len(students_list)} уч.)",
            'type': 'individual'
        })

    # Групповые уроки
    for lesson in group_lessons:
        date_str = lesson.date.strftime('%Y-%m-%d')
        if date_str not in lessons_by_date:
            lessons_by_date[date_str] = []

        students_list = [e.student.user.get_full_name() for e in lesson.enrollments.all()]
        time_str = lesson.start_time.strftime('%H:%M')

        lessons_by_date[date_str].append({
            'time': time_str,
            'subject': lesson.subject.name,
            'teacher': lesson.teacher.user.last_name,
            'students_count': len(students_list),
            'full': f"👥 {time_str} {lesson.subject.name} - Группа ({len(students_list)} уч.)",
            'type': 'group'
        })

    # Создаем сетку календаря
    calendar_days = []
    current = start_date
    week = []

    first_weekday = start_date.weekday()
    for _ in range(first_weekday):
        week.append({'day': '', 'lessons': [], 'is_current_month': False})

    while current <= end_date:
        if len(week) == 7:
            calendar_days.append(week)
            week = []

        date_str = current.strftime('%Y-%m-%d')
        week.append({
            'day': current.day,
            'lessons': lessons_by_date.get(date_str, []),
            'is_current_month': True,
            'date': current
        })
        current += timedelta(days=1)

    while len(week) < 7:
        week.append({'day': '', 'lessons': [], 'is_current_month': False})
    if week:
        calendar_days.append(week)

    # Список месяцев
    months = [
        {'value': 1, 'name': 'Январь'},
        {'value': 2, 'name': 'Февраль'},
        {'value': 3, 'name': 'Март'},
        {'value': 4, 'name': 'Апрель'},
        {'value': 5, 'name': 'Май'},
        {'value': 6, 'name': 'Июнь'},
        {'value': 7, 'name': 'Июль'},
        {'value': 8, 'name': 'Август'},
        {'value': 9, 'name': 'Сентябрь'},
        {'value': 10, 'name': 'Октябрь'},
        {'value': 11, 'name': 'Ноябрь'},
        {'value': 12, 'name': 'Декабрь'},
    ]

    # Годы
    current_year = timezone.now().year
    years = [current_year - 2, current_year - 1, current_year, current_year + 1, current_year + 2]

    context = {
        'calendar_days': calendar_days,
        'month_name': months[month - 1]['name'],
        'month': month,
        'year': year,
        'months': months,
        'years': years,
        'now': today,
        'export_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
        'lessons_count': sum(len(v) for v in lessons_by_date.values()),
    }

    html_string = render_to_string('admin/school/lesson/calendar_pdf.html', context)

    response = HttpResponse(content_type='application/pdf')
    filename = f"admin_calendar_{year}_{month:02d}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf(response)
    return response

# ============================================
# ЧАСТЬ 5: ФУНКЦИИ ИМПОРТА
# ============================================

@staff_member_required
def import_students(request):
    """Импорт учеников из Excel"""
    if request.method == 'POST':
        file = request.FILES.get('file')
        if not file:
            messages.error(request, 'Выберите файл для импорта')
            return redirect('admin:school_student_changelist')

        try:
            wb = openpyxl.load_workbook(file)
            ws = wb.active

            success_count = 0
            error_count = 0
            errors = []

            for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):
                    continue

                try:
                    student_id = row[0]
                    last_name = row[1] or ''
                    first_name = row[2] or ''
                    patronymic = row[3] or ''
                    email = row[4] or ''
                    phone = row[5] or ''
                    parent_name = row[6] or ''
                    parent_phone = row[7] or ''

                    if student_id:
                        user = User.objects.get(id=student_id)
                    else:
                        # Создаем нового с уникальным username
                        import uuid
                        username = f"student_{uuid.uuid4().hex[:8]}"
                        user = User.objects.create_user(
                            username=username,
                            email=email,
                            password='default123'
                        )
                        user.role = 'student'

                    user.last_name = last_name
                    user.first_name = first_name
                    user.patronymic = patronymic
                    user.phone = phone or ''
                    user.save()

                    student, created = Student.objects.get_or_create(user=user)
                    student.parent_name = parent_name
                    student.parent_phone = parent_phone
                    student.save()

                    success_count += 1

                except Exception as e:
                    error_count += 1
                    errors.append(f"Строка {row_num}: {str(e)}")

            if success_count > 0:
                messages.success(request, f'✅ Импортировано учеников: {success_count}')
            if error_count > 0:
                error_text = '\n'.join(errors[:5])
                if len(errors) > 5:
                    error_text += f'\n... и еще {len(errors) - 5} ошибок'
                messages.warning(request, f'⚠️ Ошибок: {error_count}\n{error_text}')

        except Exception as e:
            messages.error(request, f'Ошибка при импорте: {str(e)}')

        return redirect('admin:school_student_changelist')

    return render(request, 'admin/school/student/import.html')


@staff_member_required
def download_student_template(request):
    """Скачать шаблон для импорта учеников"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="import_students_template.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Импорт учеников"

    headers = ['ID (оставьте пустым для новых)', 'Фамилия', 'Имя', 'Отчество', 'Email', 'Телефон', 'Родитель',
               'Телефон родителя']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = openpyxl.styles.Font(bold=True)

    examples = [
        ['', 'Иванов', 'Иван', 'Иванович', 'ivanov@mail.ru', '+79001234567', 'Иванова М.И.', '+79007654321'],
        ['13', 'Петров', 'Петр', 'Петрович', 'petrov@mail.ru', '+79009876543', 'Петрова А.С.', '+79005432176'],
    ]

    for row_num, example in enumerate(examples, start=2):
        for col_num, value in enumerate(example, 1):
            ws.cell(row=row_num, column=col_num, value=value)

    wb.save(response)
    return response


@staff_member_required
def import_teachers(request):
    """Импорт учителей из Excel"""
    if request.method == 'POST':
        file = request.FILES.get('file')
        if not file:
            messages.error(request, 'Выберите файл для импорта')
            return redirect('admin:school_teacher_changelist')

        try:
            wb = openpyxl.load_workbook(file)
            ws = wb.active

            success_count = 0
            error_count = 0
            errors = []

            for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):
                    continue

                try:
                    teacher_id = row[0]
                    last_name = row[1]
                    first_name = row[2]
                    patronymic = row[3]
                    email = row[4]
                    phone = row[5]
                    subjects_str = row[6]
                    experience = row[7]
                    education = row[8]

                    if teacher_id:
                        user = User.objects.get(id=teacher_id)
                    else:
                        username = f"teacher_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                        user = User.objects.create_user(
                            username=username,
                            email=email,
                            password='default123'
                        )
                        user.role = 'teacher'

                    user.last_name = last_name
                    user.first_name = first_name
                    user.patronymic = patronymic
                    user.phone = phone
                    user.save()

                    teacher, created = Teacher.objects.get_or_create(user=user)
                    teacher.experience = experience or 0
                    teacher.education = education or ''
                    teacher.save()

                    if subjects_str:
                        subject_names = [s.strip() for s in str(subjects_str).split(';')]
                        for subject_name in subject_names:
                            subject, _ = Subject.objects.get_or_create(name=subject_name)
                            teacher.subjects.add(subject)

                    success_count += 1

                except Exception as e:
                    error_count += 1
                    errors.append(f"Строка {row_num}: {str(e)}")

            if success_count > 0:
                messages.success(request, f'✅ Импортировано учителей: {success_count}')
            if error_count > 0:
                error_text = '\n'.join(errors[:5])
                if len(errors) > 5:
                    error_text += f'\n... и еще {len(errors) - 5} ошибок'
                messages.warning(request, f'⚠️ Ошибок: {error_count}\n{error_text}')

        except Exception as e:
            messages.error(request, f'Ошибка при импорте: {str(e)}')

        return redirect('admin:school_teacher_changelist')

    return render(request, 'admin/school/teacher/import.html')


@staff_member_required
def download_teacher_template(request):
    """Скачать шаблон для импорта учителей"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="import_teachers_template.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Импорт учителей"

    headers = ['ID (пусто для новых)', 'Фамилия', 'Имя', 'Отчество', 'Email', 'Телефон', 'Предметы (через ;)',
               'Опыт (лет)', 'Образование']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = openpyxl.styles.Font(bold=True)

    examples = [
        ['', 'Соколов', 'Павел', 'Алексеевич', 'sokolov@mail.ru', '+79001112233', 'Математика;Физика', '5', 'МГУ'],
        ['10', 'Петрова', 'Анна', 'Игоревна', 'petrova@mail.ru', '+79002223344', 'Русский язык;Литература', '8',
         'МПГУ'],
    ]

    for row_num, example in enumerate(examples, start=2):
        for col_num, value in enumerate(example, 1):
            ws.cell(row=row_num, column=col_num, value=value)

    wb.save(response)
    return response


@staff_member_required
def import_lessons(request):
    """Импорт уроков из Excel или CSV"""
    if request.method == 'POST':
        file = request.FILES.get('file')
        if not file:
            messages.error(request, 'Выберите файл для импорта')
            return redirect('admin:school_lesson_changelist')

        if file.name.endswith('.csv'):
            return import_from_csv(file, request)
        elif file.name.endswith(('.xlsx', '.xls')):
            return import_from_excel(file, request)
        else:
            messages.error(request, 'Поддерживаются только файлы CSV и Excel (.xlsx, .xls)')
            return redirect('admin:school_lesson_changelist')

    return render(request, 'admin/school/lesson/import.html')


def import_from_csv(file, request):
    """Импорт из CSV"""
    try:
        decoded_file = file.read().decode('utf-8-sig').splitlines()
        reader = csv.DictReader(decoded_file, delimiter=';')

        success_count = 0
        error_count = 0
        errors = []

        for row_num, row in enumerate(reader, start=2):
            try:
                teacher_name = row.get('Учитель', '').strip()
                student_name = row.get('Ученик', '').strip()
                subject_name = row.get('Предмет', '').strip()

                teacher = find_teacher_by_full_name(teacher_name)
                if not teacher:
                    raise ValueError(f"Учитель '{teacher_name}' не найден")

                student = find_student_by_full_name(student_name)
                if not student:
                    raise ValueError(f"Ученик '{student_name}' не найден")

                subject = Subject.objects.filter(name__icontains=subject_name).first()
                if not subject:
                    raise ValueError(f"Предмет '{subject_name}' не найден")

                date_str = row.get('Дата', '').strip()
                if date_str:
                    date = datetime.strptime(date_str, '%d.%m.%Y').date()
                else:
                    raise ValueError("Дата не указана")

                start_time_str = row.get('Время начала', '').strip()
                end_time_str = row.get('Время окончания', '').strip()

                if start_time_str:
                    start_time = datetime.strptime(start_time_str, '%H:%M').time()
                else:
                    raise ValueError("Время начала не указано")

                if end_time_str:
                    end_time = datetime.strptime(end_time_str, '%H:%M').time()
                else:
                    from datetime import timedelta
                    start_dt = datetime.combine(date, start_time)
                    end_dt = start_dt + timedelta(hours=1)
                    end_time = end_dt.time()

                cost = Decimal(str(row.get('Стоимость', '1000')).replace(',', '.'))
                teacher_payment = Decimal(str(row.get('Выплата учителю', cost * Decimal('0.7'))).replace(',', '.'))

                status = row.get('Статус', 'scheduled').strip().lower()
                if status not in ['scheduled', 'completed', 'cancelled', 'overdue']:
                    status = 'scheduled'

                lesson = Lesson.objects.create(
                    teacher=teacher,
                    subject=subject,
                    date=date,
                    start_time=start_time,
                    end_time=end_time,
                    base_cost=cost,
                    base_teacher_payment=teacher_payment,
                    status=status,
                )

                LessonAttendance.objects.create(
                    lesson=lesson,
                    student=student,
                    cost=cost,
                    teacher_payment_share=teacher_payment,
                    status='registered' if status == 'scheduled' else status
                )

                success_count += 1

            except Exception as e:
                error_count += 1
                errors.append(f"Строка {row_num}: {str(e)}")

        if success_count > 0:
            messages.success(request, f'✅ Импортировано уроков: {success_count}')
        if error_count > 0:
            error_text = '\n'.join(errors[:5])
            if len(errors) > 5:
                error_text += f'\n... и еще {len(errors) - 5} ошибок'
            messages.warning(request, f'⚠️ Ошибок: {error_count}\n{error_text}')

        return redirect('admin:school_lesson_changelist')

    except Exception as e:
        messages.error(request, f'Ошибка при импорте: {str(e)}')
        return redirect('admin:school_lesson_changelist')


def import_from_excel(file, request):
    """Импорт из Excel с поддержкой ID"""
    try:
        import tempfile
        import os
        from datetime import datetime, timedelta

        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            for chunk in file.chunks():
                tmp.write(chunk)
            tmp_path = tmp.name

        wb = openpyxl.load_workbook(tmp_path)
        ws = wb.active

        headers = [cell.value for cell in ws[1] if cell.value]

        success_count = 0
        error_count = 0
        errors = []

        for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue

            try:
                row_dict = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        row_dict[header] = row[i]

                # Парсинг даты - поддерживаем оба формата
                date_str = str(row_dict.get('Дата', '')).strip()
                date = None

                # Пробуем разные форматы даты
                date_formats = [
                    '%Y-%m-%d %H:%M:%S',  # 2026-02-01 00:00:00
                    '%Y-%m-%d',  # 2026-02-01
                    '%d.%m.%Y',  # 01.02.2026
                    '%d.%m.%Y %H:%M:%S',  # 01.02.2026 00:00:00
                ]

                for fmt in date_formats:
                    try:
                        date = datetime.strptime(date_str, fmt).date()
                        break
                    except ValueError:
                        continue

                if not date:
                    raise ValueError(
                        f"Не удалось распознать дату '{date_str}'. Используйте формат ДД.ММ.ГГГГ или ГГГГ-ММ-ДД")

                # Парсинг времени
                start_time_str = str(row_dict.get('Время начала', '')).strip()
                end_time_str = str(row_dict.get('Время окончания', '')).strip()

                # Пробуем разные форматы времени
                time_formats = [
                    '%H:%M:%S',  # 10:00:00
                    '%H:%M',  # 10:00
                ]

                start_time = None
                for fmt in time_formats:
                    try:
                        start_time = datetime.strptime(start_time_str, fmt).time()
                        break
                    except ValueError:
                        continue

                if not start_time:
                    raise ValueError(f"Не удалось распознать время начала '{start_time_str}'")

                if end_time_str:
                    for fmt in time_formats:
                        try:
                            end_time = datetime.strptime(end_time_str, fmt).time()
                            break
                        except ValueError:
                            continue
                else:
                    # Если время окончания не указано, ставим +1 час
                    from datetime import timedelta, datetime
                    start_dt = datetime.combine(date, start_time)
                    end_dt = start_dt + timedelta(hours=1)
                    end_time = end_dt.time()

                # Поиск учителя по ID или ФИО
                teacher_id = row_dict.get('ID учителя')
                teacher = None

                if teacher_id:
                    teacher = find_teacher_by_id(teacher_id)
                    if not teacher:
                        raise ValueError(f"Учитель с ID '{teacher_id}' не найден")
                else:
                    teacher_name = str(row_dict.get('Учитель', '')).strip()
                    teacher = find_teacher_by_full_name(teacher_name)
                    if not teacher:
                        raise ValueError(f"Учитель '{teacher_name}' не найден")

                # Поиск учеников по ID или ФИО
                students = []

                student_ids_str = row_dict.get('ID учеников', '')
                if student_ids_str:
                    student_ids = [s.strip() for s in str(student_ids_str).split(';') if s.strip()]
                    for student_id in student_ids:
                        student = find_student_by_id(student_id)
                        if not student:
                            raise ValueError(f"Ученик с ID '{student_id}' не найден")
                        students.append(student)
                else:
                    students_str = str(row_dict.get('Ученики', '')).strip()
                    student_names = [s.strip() for s in students_str.split(';') if s.strip()]
                    for student_name in student_names:
                        student = find_student_by_full_name(student_name)
                        if not student:
                            raise ValueError(f"Ученик '{student_name}' не найден")
                        students.append(student)

                if not students:
                    raise ValueError("Не указаны ученики")

                # Поиск предмета
                subject_name = str(row_dict.get('Предмет', '')).strip()
                subject = Subject.objects.filter(name__icontains=subject_name).first()
                if not subject:
                    raise ValueError(f"Предмет '{subject_name}' не найден")

                # Стоимость
                cost_str = str(row_dict.get('Стоимость урока', '1000')).replace(',', '.')
                teacher_payment_str = str(row_dict.get('Выплата учителю', float(cost_str) * 0.7)).replace(',', '.')

                try:
                    cost = Decimal(cost_str)
                except:
                    raise ValueError(f"Неверный формат стоимости: {cost_str}")

                try:
                    teacher_payment = Decimal(teacher_payment_str)
                except:
                    teacher_payment = cost * Decimal('0.7')

                # Статус
                status = str(row_dict.get('Статус', 'scheduled')).strip().lower()
                if status not in ['scheduled', 'completed', 'cancelled', 'overdue']:
                    status = 'scheduled'

                # Создание урока
                lesson = Lesson.objects.create(
                    teacher=teacher,
                    subject=subject,
                    date=date,
                    start_time=start_time,
                    end_time=end_time,
                    base_cost=cost,
                    base_teacher_payment=teacher_payment,
                    status=status,
                )

                # Создаем записи посещаемости для всех учеников
                for student in students:
                    LessonAttendance.objects.create(
                        lesson=lesson,
                        student=student,
                        cost=cost,
                        teacher_payment_share=teacher_payment,
                        status='registered' if status == 'scheduled' else status
                    )

                success_count += 1

            except Exception as e:
                error_count += 1
                errors.append(f"Строка {row_num}: {str(e)}")

        # Удаляем временный файл
        os.unlink(tmp_path)

        # Сообщаем результат
        if success_count > 0:
            messages.success(request, f'✅ Импортировано уроков: {success_count}')
        if error_count > 0:
            error_text = '\n'.join(errors[:5])
            if len(errors) > 5:
                error_text += f'\n... и еще {len(errors) - 5} ошибок'
            messages.warning(request, f'⚠️ Ошибок: {error_count}\n{error_text}')

        return redirect('admin:school_lesson_changelist')

    except Exception as e:
        messages.error(request, f'Ошибка при импорте: {str(e)}')
        return redirect('admin:school_lesson_changelist')


@staff_member_required
def download_user_template(request):
    """Скачать шаблон для импорта пользователей"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from django.http import HttpResponse
    from datetime import datetime

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Импорт пользователей"

    headers = ['Username', 'Имя', 'Фамилия', 'Отчество', 'Email', 'Телефон', 'Роль', 'Пароль']

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="417690", end_color="417690", fill_type="solid")

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')

    # Пример данных
    examples = [
        ['ivanov', 'Иван', 'Иванов', 'Иванович', 'ivan@mail.ru', '+79991234567', 'student', 'pass123'],
        ['petrova', 'Мария', 'Петрова', 'Сергеевна', 'maria@mail.ru', '+79997654321', 'teacher', 'pass123'],
    ]

    for row_num, example in enumerate(examples, start=2):
        for col_num, value in enumerate(example, 1):
            ws.cell(row=row_num, column=col_num, value=value)

    column_widths = [15, 15, 15, 15, 25, 15, 10, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f"user_import_template_{datetime.now().strftime('%Y%m%d')}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    wb.save(response)
    return response


from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import render, redirect
from django.contrib import messages
import openpyxl
from datetime import datetime
import traceback
from .models import User


@staff_member_required
def import_users_view(request):
    """Отдельное представление для импорта пользователей"""
    if request.method == 'POST' and request.FILES.get('import_file'):
        file = request.FILES['import_file']

        try:
            wb = openpyxl.load_workbook(file)
            ws = wb.active

            success_count = 0
            error_count = 0
            errors = []

            for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):
                    continue

                try:
                    username = str(row[0]) if row[0] else None
                    first_name = str(row[1]) if row[1] else ''
                    last_name = str(row[2]) if row[2] else ''
                    patronymic = str(row[3]) if row[3] else ''
                    email = str(row[4]) if row[4] else ''
                    phone = str(row[5]) if row[5] else ''
                    role = str(row[6]) if row[6] else 'student'
                    password = str(row[7]) if row[7] else 'default123'

                    if not username:
                        raise ValueError("Имя пользователя обязательно")

                    if User.objects.filter(username=username).exists():
                        raise ValueError(f"Пользователь с username '{username}' уже существует")

                    if email and User.objects.filter(email=email).exists():
                        raise ValueError(f"Пользователь с email '{email}' уже существует")

                    user = User.objects.create_user(
                        username=username,
                        password=password,
                        first_name=first_name,
                        last_name=last_name,
                        email=email,
                        phone=phone,
                        role=role
                    )
                    user.patronymic = patronymic
                    user.save()

                    success_count += 1

                except Exception as e:
                    error_count += 1
                    errors.append(f"Строка {row_num}: {str(e)}")

            messages.success(request, f'✅ Импортировано пользователей: {success_count}')
            if error_count > 0:
                error_text = '\n'.join(errors[:5])
                if len(errors) > 5:
                    error_text += f'\n... и еще {len(errors) - 5} ошибок'
                messages.warning(request, f'⚠️ Ошибок: {error_count}\n{error_text}')

        except Exception as e:
            messages.error(request, f'Ошибка при импорте: {str(e)}')

        return redirect('admin:school_user_changelist')

    return render(request, 'admin/school/user/import.html')


# ============================================
# ЧАСТЬ 6: API И JSON ФУНКЦИИ
# ============================================

@require_GET
def api_schedules(request):
    """API для календаря расписаний"""
    schedules = Schedule.objects.filter(is_active=True).select_related('teacher__user')

    events = []
    today = date.today()

    for schedule in schedules:
        for i in range(30):
            event_date = today + timedelta(days=i)
            if event_date.weekday() == schedule.day_of_week:
                lesson = Lesson.objects.filter(
                    teacher=schedule.teacher,
                    date=event_date,
                    start_time=schedule.start_time
                ).first()

                start_dt = datetime.combine(event_date, schedule.start_time)
                end_dt = datetime.combine(event_date, schedule.end_time)

                event = {
                    'id': f"schedule_{schedule.id}_{event_date}",
                    'teacher_name': schedule.teacher.user.get_full_name(),
                    'subject': 'Расписание',
                    'start': start_dt.isoformat(),
                    'end': end_dt.isoformat(),
                    'color': '#28a745' if lesson else '#3788d8',
                }

                if lesson:
                    event['subject'] = lesson.subject.name
                    event['student_name'] = lesson.student.user.get_full_name()
                    event['status'] = lesson.status

                events.append(event)

    return JsonResponse(events, safe=False)


@staff_member_required
def schedule_calendar_data(request):
    """API для календаря расписаний"""
    schedules = Schedule.objects.filter(is_active=True).select_related('teacher__user')

    events = []
    today = date.today()

    for i in range(60):
        event_date = today + timedelta(days=i)
        day_schedules = schedules.filter(date=event_date)

        for schedule in day_schedules:
            lesson = Lesson.objects.filter(
                teacher=schedule.teacher,
                date=event_date,
                start_time=schedule.start_time
            ).first()

            start_dt = datetime.combine(event_date, schedule.start_time)
            end_dt = datetime.combine(event_date, schedule.end_time)

            color = '#79aec8'
            if lesson:
                if lesson.status == 'completed':
                    color = '#28a745'
                elif lesson.status == 'overdue':
                    color = '#dc3545'
                elif lesson.status == 'scheduled':
                    color = '#007bff'
                elif lesson.status == 'cancelled':
                    color = '#fd7e14'

            event = {
                'id': f"schedule_{schedule.id}_{event_date}",
                'schedule_id': schedule.id,
                'teacher_name': schedule.teacher.user.get_full_name(),
                'start': start_dt.isoformat(),
                'end': end_dt.isoformat(),
                'color': color,
            }

            if lesson:
                event['lesson_id'] = lesson.id
                event['title'] = f"{schedule.teacher.user.last_name} - {lesson.subject.name}"
                first_attendance = lesson.attendance.first()
                if first_attendance:
                    event['title'] += f" ({first_attendance.student.user.last_name})"
            else:
                event['title'] = f"{schedule.teacher.user.last_name} - свободно"

            events.append(event)

    return JsonResponse(events, safe=False)


@login_required
def get_notifications(request):
    """API для получения уведомлений"""
    try:
        notifications = Notification.objects.filter(user=request.user).order_by('-created_at')[:20]
        unread_count = Notification.objects.filter(user=request.user, is_read=False).count()

        notifications_data = []
        for n in notifications:
            try:
                time_diff = timezone.now() - n.created_at
                if time_diff.days > 0:
                    created_ago = f"{time_diff.days} дн. назад"
                elif time_diff.seconds // 3600 > 0:
                    created_ago = f"{time_diff.seconds // 3600} ч. назад"
                elif time_diff.seconds // 60 > 0:
                    created_ago = f"{time_diff.seconds // 60} мин. назад"
                else:
                    created_ago = "только что"
            except:
                created_ago = n.created_at.strftime('%d.%m.%Y %H:%M')

            notifications_data.append({
                'id': n.id,
                'title': n.title,
                'message': n.message,
                'type': n.notification_type,
                'is_read': n.is_read,
                'link': n.link if n.link else '',
                'created_at': n.created_at.strftime('%d.%m.%Y %H:%M'),
                'created_ago': created_ago,
            })

        return JsonResponse({
            'unread_count': unread_count,
            'notifications': notifications_data
        })

    except Exception as e:
        print(f"❌ Ошибка в get_notifications: {e}")
        return JsonResponse({'error': str(e), 'notifications': [], 'unread_count': 0}, status=500)


@login_required
@require_POST
def mark_notification_read(request, notification_id):
    """Отметить уведомление как прочитанное"""
    try:
        notification = Notification.objects.get(id=notification_id, user=request.user)
        notification.is_read = True
        notification.save()

        unread_count = Notification.objects.filter(user=request.user, is_read=False).count()

        return JsonResponse({
            'status': 'ok',
            'unread_count': unread_count,
            'message': 'Уведомление отмечено как прочитанное'
        })
    except Notification.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Уведомление не найдено'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
@require_POST
def mark_all_notifications_read(request):
    """Отметить все уведомления как прочитанные"""
    try:
        count = Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
        return JsonResponse({
            'status': 'ok',
            'count': count,
            'message': f'Отмечено {count} уведомлений'
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


def generate_video_room(request, lesson_id):
    """Генерирует комнату для видео"""
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if not lesson.video_room:
        lesson.video_room = str(uuid.uuid4())[:8]
        lesson.save()

    return JsonResponse({
        'room': lesson.video_room,
        'url': f'https://meet.jit.si/plusprogress-{lesson.id}-{lesson.date}'
    })


@login_required
@require_POST
def create_video_room(request, lesson_id):
    """Учитель создает видео-комнату для урока"""
    try:
        lesson = get_object_or_404(Lesson, id=lesson_id)

        # Проверка доступа
        if request.user.role != 'teacher' or lesson.teacher.user != request.user:
            return JsonResponse({'error': 'Доступ запрещен'}, status=403)

        if lesson.status != 'scheduled':
            return JsonResponse({'error': 'Урок уже проведен или отменен'}, status=400)

        # ✅ ТОЛЬКО ЭТО НОВОЕ - логирование видео
        log_user_action(
            request,
            'video_room_enter',
            f'Вход в видео-комнату урока #{lesson.id}',
            object_id=lesson.id,
            object_type='lesson'
        )

        if not lesson.video_room:
            import uuid
            lesson.video_room = str(uuid.uuid4())[:8]
            lesson.save()

        return JsonResponse({
            'success': True,
            'room': lesson.video_room
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
# ============================================
# ЧАСТЬ 7: ОТЧЕТЫ
# ============================================

@login_required
def overdue_report(request):
    """Отчет по просроченным занятиям"""
    if request.user.role not in ['admin', 'teacher']:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    now = timezone.now()
    overdue_lessons = Lesson.objects.filter(
        status='scheduled',
        date__lt=now.date()
    ) | Lesson.objects.filter(
        status='scheduled',
        date=now.date(),
        start_time__lt=now.time()
    )

    for lesson in overdue_lessons:
        lesson.check_overdue()

    teacher_id = request.GET.get('teacher')
    student_id = request.GET.get('student')
    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')

    lessons = Lesson.objects.filter(status='overdue').select_related(
        'teacher__user', 'subject'
    ).prefetch_related('attendance__student__user')

    if teacher_id:
        lessons = lessons.filter(teacher_id=teacher_id)
    if student_id:
        lessons = lessons.filter(attendance__student_id=student_id)
    if date_from:
        lessons = lessons.filter(date__gte=date_from)
    if date_to:
        lessons = lessons.filter(date__lte=date_to)

    # ИСПОЛЬЗУЕМ PeriodFinanceCalculator для статистики
    period_calc = PeriodFinanceCalculator(lessons)
    stats = period_calc.lessons_stats

    context = {
        'lessons': lessons.order_by('-date', '-start_time'),
        'stats': {
            'total': stats['total'],
            'by_teacher': lessons.values('teacher__user__last_name').annotate(count=Count('id')).order_by('-count'),
            'by_subject': lessons.values('subject__name').annotate(count=Count('id')).order_by('-count'),
        },
        'teachers': Teacher.objects.all(),
        'students': Student.objects.all(),
    }

    return render(request, 'school/reports/overdue.html', context)


@staff_member_required
def student_report(request, student_id):
    """Отчет по ученику"""
    from django.db.models import Sum
    student = get_object_or_404(Student, id=student_id)

    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')

    # Преобразуем строки в даты
    start_date = None
    end_date = None
    if date_from:
        start_date = datetime.strptime(date_from, '%Y-%m-%d').date()
    if date_to:
        end_date = datetime.strptime(date_to, '%Y-%m-%d').date()

    # Получаем ТОЛЬКО ПРОВЕДЕННЫЕ уроки (attended)
    attendances = LessonAttendance.objects.filter(
        student=student,
        status='attended'
    ).select_related(
        'lesson', 'lesson__subject'
    ).order_by('lesson__date')

    if start_date:
        attendances = attendances.filter(lesson__date__gte=start_date)
    if end_date:
        attendances = attendances.filter(lesson__date__lte=end_date)

    # Получаем все уникальные даты
    dates = attendances.dates('lesson__date', 'day').order_by('lesson__date')

    # Получаем все уникальные предметы
    subjects = attendances.values_list('lesson__subject__name', flat=True).distinct()

    # Создаем словарь для хранения данных по предметам
    subjects_data_dict = {}
    daily_totals = {date: 0 for date in dates}

    # Инициализируем словарь для каждого предмета
    for subject_name in subjects:
        subjects_data_dict[subject_name] = {date: 0 for date in dates}

    # Заполняем данные
    for attendance in attendances:
        subject_name = attendance.lesson.subject.name
        lesson_date = attendance.lesson.date
        cost = attendance.cost

        subjects_data_dict[subject_name][lesson_date] += cost
        daily_totals[lesson_date] += cost

    # Формируем данные для таблицы
    subjects_data = []
    total_sum = 0

    for subject_name, daily_costs in subjects_data_dict.items():
        daily_costs_list = []
        subject_total = 0

        for date in dates:
            cost = daily_costs.get(date, 0)
            daily_costs_list.append(float(cost))
            subject_total += cost

        subjects_data.append({
            'name': subject_name,
            'daily_costs': daily_costs_list,
            'total': float(subject_total)
        })
        total_sum += subject_total

    # Статистика по ученику
    total_lessons = attendances.count()
    total_attended_cost = attendances.aggregate(Sum('cost'))['cost__sum'] or 0

    # Получаем отдельно уроки в долг (для статистики)
    debt_attendances = LessonAttendance.objects.filter(
        student=student,
        status='debt'
    )
    if start_date:
        debt_attendances = debt_attendances.filter(lesson__date__gte=start_date)
    if end_date:
        debt_attendances = debt_attendances.filter(lesson__date__lte=end_date)

    debt_lessons = debt_attendances.count()
    total_debt_cost = debt_attendances.aggregate(Sum('cost'))['cost__sum'] or 0

    # ===== ИСПОЛЬЗУЕМ НОВЫЙ МЕТОД =====
    # Пополнения счета (income)
    total_deposits = Payment.objects.filter(
        user=student.user,
        payment_type='income'
    ).aggregate(Sum('amount'))['amount__sum'] or 0

    # Правильный баланс через метод пользователя
    correct_balance = student.user.get_balance()

    context = {
        'student': student,
        'dates': dates,
        'subjects_data': subjects_data,
        'daily_totals': [float(daily_totals.get(date, 0)) for date in dates],
        'total_lessons': total_lessons,
        'total_attended_cost': float(total_attended_cost),
        'debt_lessons': debt_lessons,
        'total_debt_cost': float(total_debt_cost),
        'student_balance': correct_balance,  # ✅ Правильный баланс
        'total_deposits': float(total_deposits),
        'total_expenses': float(Payment.objects.filter(
            user=student.user,
            payment_type='expense'
        ).aggregate(Sum('amount'))['amount__sum'] or 0),
    }

    return render(request, 'admin/school/student/report.html', context)

@staff_member_required
def teacher_report(request, teacher_id):
    """Отчет по учителю"""
    teacher = get_object_or_404(Teacher, id=teacher_id)

    date_from = request.GET.get('date_from')
    date_to = request.GET.get('date_to')

    # Преобразуем строки в даты
    start_date = None
    end_date = None
    if date_from:
        start_date = datetime.strptime(date_from, '%Y-%m-%d').date()
    if date_to:
        end_date = datetime.strptime(date_to, '%Y-%m-%d').date()

    # ОТЛАДКА: смотрим все статусы уроков
    all_statuses = Lesson.objects.filter(teacher=teacher).values_list('status', flat=True).distinct()
    print(f"\n🔍 Все статусы уроков учителя: {list(all_statuses)}")

    # Получаем ТОЛЬКО ПРОВЕДЕННЫЕ уроки
    lessons = Lesson.objects.filter(
        teacher=teacher,
        status='completed'
    ).prefetch_related(
        'attendance__student__user', 'subject'
    ).order_by('date')

    # ОТЛАДКА: сколько нашлось
    print(f"🔍 Найдено проведенных уроков: {lessons.count()}")

    if start_date:
        lessons = lessons.filter(date__gte=start_date)
    if end_date:
        lessons = lessons.filter(date__lte=end_date)

    # Получаем все уникальные даты
    dates = lessons.dates('date', 'day').order_by('date')

    # ОТЛАДКА: даты
    print(f"🔍 Даты проведенных уроков: {[d.strftime('%d.%m.%Y') for d in dates]}")

    # Словари для хранения данных
    students_lessons_dict = {}  # Для стоимости уроков (ученик платит)
    students_earnings_dict = {}  # Для заработка учителя (teacher_payment_share)
    daily_totals_lessons = {date: 0 for date in dates}
    daily_totals_earnings = {date: 0 for date in dates}

    total_lessons_count = 0
    total_income_sum = 0
    total_earnings_sum = 0

    # Собираем данные по каждому уроку
    for lesson in lessons:
        total_lessons_count += 1

        for attendance in lesson.attendance.all():
            student_name = attendance.student.user.get_full_name()
            subject_name = lesson.subject.name
            key = f"{student_name} ({subject_name})"

            # Стоимость для ученика
            cost = attendance.cost
            # Заработок учителя
            earning = attendance.teacher_payment_share

            # Добавляем в словарь стоимости (ученик платит)
            if key not in students_lessons_dict:
                students_lessons_dict[key] = {date: 0 for date in dates}
            students_lessons_dict[key][lesson.date] += cost

            # Добавляем в словарь заработка учителя
            if key not in students_earnings_dict:
                students_earnings_dict[key] = {date: 0 for date in dates}
            students_earnings_dict[key][lesson.date] += earning

            # Обновляем итоги по дням
            daily_totals_lessons[lesson.date] += cost
            daily_totals_earnings[lesson.date] += earning

            # Обновляем общие итоги
            total_income_sum += cost
            total_earnings_sum += earning

    # Формируем данные для таблиц
    lessons_data = []
    earnings_data = []

    for key in students_lessons_dict.keys():
        # Данные по стоимости уроков
        daily_costs = []
        student_total = 0
        for date in dates:
            cost = students_lessons_dict[key].get(date, 0)
            daily_costs.append(float(cost))
            student_total += cost

        lessons_data.append({
            'name': key,
            'daily_costs': daily_costs,
            'total': float(student_total)
        })

        # Данные по заработку учителя
        daily_earnings = []
        earning_total = 0
        for date in dates:
            earning = students_earnings_dict[key].get(date, 0)
            daily_earnings.append(float(earning))
            earning_total += earning

        earnings_data.append({
            'name': key,
            'daily_earnings': daily_earnings,
            'total': float(earning_total)
        })

    # Сортируем данные по имени ученика
    lessons_data.sort(key=lambda x: x['name'])
    earnings_data.sort(key=lambda x: x['name'])

    # Формируем итоги по дням
    daily_totals_lessons_list = []
    daily_totals_earnings_list = []

    for date in dates:
        daily_totals_lessons_list.append(float(daily_totals_lessons.get(date, 0)))
        daily_totals_earnings_list.append(float(daily_totals_earnings.get(date, 0)))

    context = {
        'teacher': teacher,
        'dates': dates,
        'lessons_data': lessons_data,
        'earnings_data': earnings_data,
        'daily_totals_lessons': daily_totals_lessons_list,
        'daily_totals_earnings': daily_totals_earnings_list,
        'total_lessons': total_lessons_count,
        'total_income': float(total_income_sum),
        'total_earnings': float(total_earnings_sum),
    }

    # ИТОГОВАЯ ОТЛАДКА
    print(f"\n{'=' * 60}")
    print(f"ОТЧЕТ ПО УЧИТЕЛЮ: {teacher.user.get_full_name()}")
    print(f"Всего проведенных уроков в отчете: {total_lessons_count}")
    print(f"Общая стоимость: {total_income_sum}")
    print(f"Заработок: {total_earnings_sum}")
    print(f"Даты в отчете: {[d.strftime('%d.%m.%Y') for d in dates]}")
    print(f"{'=' * 60}\n")

    return render(request, 'admin/school/teacher/report.html', context)


@staff_member_required
def teacher_payments_dashboard(request):
    """Дашборд для расчета выплат учителям"""
    teachers = Teacher.objects.all().select_related('user')

    end_date = datetime.now().date()
    start_date = end_date - timedelta(days=7)

    context = {
        'teachers': teachers,
        'default_start': start_date.strftime('%Y-%m-%d'),
        'default_end': end_date.strftime('%Y-%m-%d'),
        'title': 'Расчет выплат учителям',
    }
    return render(request, 'admin/school/teacher_payments/dashboard.html', context)


@staff_member_required
def calculate_teacher_payment(request):
    """API для расчета выплат учителю за период"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не поддерживается'}, status=405)

    try:
        data = json.loads(request.body)
        teacher_id = data.get('teacher_id')
        start_date = datetime.strptime(data.get('start_date'), '%Y-%m-%d').date()
        end_date = datetime.strptime(data.get('end_date'), '%Y-%m-%d').date()

        teacher = get_object_or_404(Teacher, id=teacher_id)

        completed_lessons = Lesson.objects.filter(
            teacher=teacher,
            status='completed',
            date__gte=start_date,
            date__lte=end_date
        ).prefetch_related('attendance__student__user', 'subject')

        # ИСПОЛЬЗУЕМ PeriodFinanceCalculator
        period_calc = PeriodFinanceCalculator(completed_lessons)
        stats = period_calc.lessons_stats

        # Агрегация по предметам
        subject_stats = []
        subject_totals = {}

        for lesson in completed_lessons:
            calculator = LessonFinanceCalculator(lesson)
            subject_name = lesson.subject.name
            if subject_name not in subject_totals:
                subject_totals[subject_name] = {'count': 0, 'payment': 0}
            for attendance in calculator.get_attendance_details():
                subject_totals[subject_name]['count'] += 1
                subject_totals[subject_name]['payment'] += attendance['teacher_payment']

        for name, data in subject_totals.items():
            subject_stats.append({
                'subject__name': name,
                'lesson_count': data['count'],
                'total_payment': data['payment']
            })

        # Агрегация по ученикам
        student_stats = []
        student_totals = {}

        for lesson in completed_lessons:
            calculator = LessonFinanceCalculator(lesson)
            for attendance in calculator.get_attendance_details():
                student_name = attendance['student_name']
                if student_name not in student_totals:
                    student_totals[student_name] = {'count': 0, 'payment': 0}
                student_totals[student_name]['count'] += 1
                student_totals[student_name]['payment'] += attendance['teacher_payment']

        for name, data in student_totals.items():
            name_parts = name.split()
            student_stats.append({
                'student__user__last_name': name_parts[0] if name_parts else '',
                'student__user__first_name': name_parts[1] if len(name_parts) > 1 else '',
                'student__user__patronymic': '',
                'lesson_count': data['count'],
                'total_payment': data['payment']
            })

        # Данные для таблицы по дням
        dates = []
        current_date = start_date
        while current_date <= end_date:
            dates.append(current_date.strftime('%d.%m.%Y'))
            current_date += timedelta(days=1)

        lessons_data = []
        for lesson in completed_lessons:
            calculator = LessonFinanceCalculator(lesson)
            for attendance in calculator.get_attendance_details():
                lessons_data.append({
                    'date': lesson.date.strftime('%d.%m.%Y'),
                    'student': attendance['student_name'],
                    'subject': lesson.subject.name,
                    'cost': attendance['cost'],
                    'teacher_payment': attendance['teacher_payment'],
                    'status': lesson.status
                })

        response_data = {
            'teacher': {
                'id': teacher.id,
                'name': teacher.user.get_full_name(),
            },
            'period': {
                'start': start_date.strftime('%d.%m.%Y'),
                'end': end_date.strftime('%d.%m.%Y'),
            },
            'totals': {
                'lessons': stats['completed'],
                'cost': stats['total_cost'],
                'payment': stats['teacher_payment'],
            },
            'subject_stats': subject_stats,
            'student_stats': sorted(student_stats, key=lambda x: x['total_payment'], reverse=True),
            'lessons_data': lessons_data,
            'dates': dates,
        }

        return JsonResponse(response_data)

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


# ============================================
# ЧАСТЬ 8: ДОМАШНИЕ ЗАДАНИЯ
# ============================================

@login_required
def teacher_homeworks(request):
    """Список заданий для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    student_id = request.GET.get('student')
    status = request.GET.get('status')

    homeworks = Homework.objects.filter(teacher=teacher).select_related(
        'student__user', 'subject'
    ).prefetch_related('submission')

    if student_id:
        homeworks = homeworks.filter(student_id=student_id)

    students = Student.objects.filter(teachers=teacher)

    stats = {
        'total': homeworks.count(),
        'pending': sum(1 for h in homeworks if h.get_status() == 'pending'),
        'submitted': sum(1 for h in homeworks if h.get_status() == 'submitted'),
        'checked': sum(1 for h in homeworks if h.get_status() == 'checked'),
        'overdue': sum(1 for h in homeworks if h.get_status() == 'overdue'),
    }

    context = {
        'homeworks': homeworks.order_by('-created_at'),
        'students': students,
        'stats': stats,
        'teacher': teacher,
    }
    return render(request, 'school/teacher/homeworks.html', context)


@login_required
# school/views.py - ЗАМЕНИТЕ существующую функцию

@login_required
def teacher_homework_create(request, student_id=None):
    """Создание домашнего задания (из дашборда или для конкретного ученика)"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    # Если передан student_id, проверяем что ученик принадлежит учителю
    student = None
    if student_id:
        student = get_object_or_404(Student, id=student_id, teachers=teacher)

    if request.method == 'POST':
        # Получаем данные из формы
        student_id = request.POST.get('student')
        subject_id = request.POST.get('subject')
        lesson_id = request.POST.get('lesson')  # может быть пустым
        title = request.POST.get('title')
        description = request.POST.get('description')
        deadline = request.POST.get('deadline')
        attachments = request.FILES.get('attachments')

        # Валидация
        if not all([student_id, subject_id, title, description, deadline]):
            messages.error(request, 'Заполните все обязательные поля')
            return redirect('teacher_homework_create')

        try:
            student = get_object_or_404(Student, id=student_id, teachers=teacher)
            subject = get_object_or_404(Subject, id=subject_id)

            # Преобразуем строку даты в datetime
            deadline_dt = datetime.fromisoformat(deadline)
            if timezone.is_naive(deadline_dt):
                deadline_dt = timezone.make_aware(deadline_dt)

            if deadline_dt < timezone.now():
                messages.error(request, 'Срок сдачи не может быть в прошлом')
                return redirect('teacher_homework_create')

            # Создаем ДЗ
            homework_data = {
                'teacher': teacher,
                'student': student,
                'subject': subject,
                'title': title,
                'description': description,
                'deadline': deadline_dt,
                'attachments': attachments,
            }

            # Если выбран урок, добавляем его
            if lesson_id:
                lesson = get_object_or_404(Lesson, id=lesson_id, teacher=teacher)
                homework_data['lesson'] = lesson

            homework = Homework.objects.create(**homework_data)

            # ✅ ЛОГИРОВАНИЕ создания ДЗ
            additional_data = {
                'student': student.user.get_full_name(),
                'subject': subject.name,
                'title': title,
                'deadline': deadline
            }
            if lesson_id:
                additional_data['lesson_id'] = lesson_id

            log_user_action(
                request,
                'homework_create',
                f'Создано домашнее задание #{homework.id} - {title} для {student.user.get_full_name()}',
                object_id=homework.id,
                object_type='homework',
                additional_data=additional_data
            )

            # ✅ ВНУТРЕННЕЕ УВЕДОМЛЕНИЕ ученику
            Notification.objects.create(
                user=student.user,
                title='📝 Новое домашнее задание',
                message=f'Учитель {teacher.user.get_full_name()} выдал задание: {title} по предмету {subject.name}. Срок сдачи: {deadline_dt.strftime("%d.%m.%Y %H:%M")}',
                notification_type='homework_assigned',
                link=f'/student/homework/{homework.id}/'
            )

            # ✅ ВНУТРЕННЕЕ УВЕДОМЛЕНИЕ учителю (подтверждение)
            Notification.objects.create(
                user=request.user,
                title='✅ Домашнее задание создано',
                message=f'Задание "{title}" для {student.user.get_full_name()} успешно создано',
                notification_type='system',
                link=f'/teacher/homework/{homework.id}/'
            )

            # ✅ УВЕДОМЛЕНИЕ В TELEGRAM
            try:
                from school.telegram import notify_new_homework
                notify_new_homework(homework)
            except Exception as e:
                print(f"❌ Ошибка отправки Telegram: {e}")

            messages.success(request, f'✅ Домашнее задание "{title}" успешно создано')
            return redirect('teacher_homework_detail', homework_id=homework.id)

        except Exception as e:
            print(f"❌ Ошибка при создании ДЗ: {e}")
            import traceback
            traceback.print_exc()
            messages.error(request, f'Ошибка при создании задания: {str(e)}')
            return redirect('teacher_homework_create')

    # GET запрос - показываем форму
    students = teacher.student_set.all().select_related('user')
    subjects = teacher.subjects.all()

    # Текущая дата + 7 дней для дедлайна по умолчанию
    default_deadline = (timezone.now() + timedelta(days=7)).strftime('%Y-%m-%dT%H:%M')

    context = {
        'teacher': teacher,
        'students': students,
        'subjects': subjects,
        'default_deadline': default_deadline,
        'selected_student': student,
    }
    return render(request, 'school/teacher/homework_create.html', context)


@login_required
def teacher_homework_detail(request, homework_id):
    """Детали задания для учителя с возможностью проверки"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    homework = get_object_or_404(Homework, id=homework_id, teacher=teacher)

    # ✅ ЛОГИРОВАНИЕ просмотра
    log_user_action(
        request,
        'homework_view',
        f'Просмотр домашнего задания #{homework.id} - {homework.title}',
        object_id=homework.id,
        object_type='homework',
        additional_data={
            'student': homework.student.user.get_full_name(),
            'subject': homework.subject.name
        }
    )

    submission = None
    if hasattr(homework, 'submission'):
        submission = homework.submission

    if request.method == 'POST' and submission:
        form = HomeworkCheckForm(request.POST, instance=submission)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.status = 'checked'
            submission.checked_at = timezone.now()
            submission.save()

            # ✅ ЛОГИРОВАНИЕ проверки
            log_user_action(
                request,
                'homework_check',
                f'Проверено домашнее задание #{homework.id}, оценка: {submission.grade}',
                object_id=homework.id,
                object_type='homework',
                additional_data={
                    'grade': submission.grade,
                    'comment': submission.teacher_comment,
                    'student': homework.student.user.get_full_name()
                }
            )

            # ✅ ВНУТРЕННЕЕ УВЕДОМЛЕНИЕ ученику
            Notification.objects.create(
                user=homework.student.user,
                title='✅ Задание проверено',
                message=f"Ваше задание '{homework.title}' проверено. Оценка: {submission.grade}",
                notification_type='homework_checked',
                link=f'/student/homework/{homework.id}/'
            )

            # ✅ УВЕДОМЛЕНИЕ В TELEGRAM
            try:
                from school.telegram import notify_homework_checked
                notify_homework_checked(homework, submission)
            except Exception as e:
                print(f"❌ Ошибка отправки Telegram: {e}")

            messages.success(request, '✅ Задание проверено')
            return redirect('teacher_homework_detail', homework_id=homework.id)
    else:
        form = HomeworkCheckForm(instance=submission) if submission else None

    context = {
        'homework': homework,
        'submission': submission,
        'form': form,
        'teacher': teacher,
    }
    return render(request, 'school/teacher/homework_detail.html', context)

@login_required
def student_homeworks(request):
    """Список заданий для ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    student = request.user.student_profile
    status = request.GET.get('status')

    homeworks = Homework.objects.filter(student=student).select_related(
        'teacher__user', 'subject'
    ).prefetch_related('submission')

    if status:
        if status == 'pending':
            homeworks = [h for h in homeworks if h.get_status() == 'pending']
        elif status == 'submitted':
            homeworks = [h for h in homeworks if h.get_status() == 'submitted']
        elif status == 'checked':
            homeworks = [h for h in homeworks if h.get_status() == 'checked']
        elif status == 'overdue':
            homeworks = [h for h in homeworks if h.get_status() == 'overdue']

    all_homeworks = Homework.objects.filter(student=student)
    stats = {
        'total': all_homeworks.count(),
        'pending': sum(1 for h in all_homeworks if h.get_status() == 'pending'),
        'submitted': sum(1 for h in all_homeworks if h.get_status() == 'submitted'),
        'checked': sum(1 for h in all_homeworks if h.get_status() == 'checked'),
        'overdue': sum(1 for h in all_homeworks if h.get_status() == 'overdue'),
    }

    context = {
        'homeworks': homeworks,
        'stats': stats,
        'student': student,
    }
    return render(request, 'school/student/homeworks.html', context)


@login_required
def student_homework_detail(request, homework_id):
    """Детали задания для ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    student = request.user.student_profile
    homework = get_object_or_404(Homework, id=homework_id, student=student)

    try:
        submission = homework.submission
        can_submit = False
    except HomeworkSubmission.DoesNotExist:
        submission = None
        can_submit = True

    if request.method == 'POST' and can_submit:
        form = HomeworkSubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.homework = homework
            submission.student = student
            submission.save()

            messages.success(request, 'Задание отправлено на проверку!')
            return redirect('student_homeworks')
    else:
        form = HomeworkSubmissionForm()

    context = {
        'homework': homework,
        'submission': submission,
        'form': form if can_submit else None,
        'can_submit': can_submit,
        'student': student,
    }
    return render(request, 'school/student/homework_detail.html', context)


# ============================================
# ЧАСТЬ 9: ОТЗЫВЫ И ОЦЕНКИ
# ============================================

@login_required
def lesson_feedback(request, lesson_id):
    """Страница оценки урока"""
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    try:
        attendance = lesson.attendance.get(student__user=request.user)
    except LessonAttendance.DoesNotExist:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    if lesson.status != 'completed':
        messages.error(request, 'Можно оценивать только проведенные уроки')
        return redirect('student_dashboard')

    if hasattr(lesson, 'feedback'):
        messages.info(request, 'Вы уже оценили этот урок')
        return redirect('student_dashboard')

    if request.method == 'POST':
        form = LessonFeedbackForm(request.POST)
        if form.is_valid():
            feedback = form.save(commit=False)
            feedback.lesson = lesson
            feedback.student = attendance.student
            feedback.teacher = lesson.teacher
            feedback.save()

            messages.success(request, 'Спасибо за вашу оценку! Отзыв поможет нам стать лучше.')
            return redirect('student_dashboard')
    else:
        form = LessonFeedbackForm()

    context = {
        'lesson': lesson,
        'form': form,
    }
    return render(request, 'school/student/lesson_feedback.html', context)


@login_required
def teacher_feedbacks(request):
    """Страница с отзывами для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    feedbacks = LessonFeedback.objects.filter(teacher=teacher).select_related(
        'lesson', 'student__user', 'lesson__subject'
    ).order_by('-created_at')

    stats = feedbacks.aggregate(
        avg_rating=Avg('rating'),
        total=Count('id')
    )

    rating_distribution = {
        5: feedbacks.filter(rating=5).count(),
        4: feedbacks.filter(rating=4).count(),
        3: feedbacks.filter(rating=3).count(),
        2: feedbacks.filter(rating=2).count(),
        1: feedbacks.filter(rating=1).count(),
    }

    context = {
        'feedbacks': feedbacks,
        'stats': stats,
        'rating_distribution': rating_distribution,
        'teacher': teacher,
    }
    return render(request, 'school/teacher/feedbacks.html', context)


@login_required
def student_feedbacks(request):
    """Страница с отзывами для ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    student = request.user.student_profile
    feedbacks = LessonFeedback.objects.filter(student=student).select_related(
        'lesson', 'teacher__user', 'lesson__subject'
    ).order_by('-created_at')

    context = {
        'feedbacks': feedbacks,
        'student': student,
    }
    return render(request, 'school/student/feedbacks.html', context)


# ============================================
# ЧАСТЬ 10: ГРУППОВЫЕ УРОКИ
# ============================================

@login_required
def teacher_group_lessons(request):
    """Список групповых уроков для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    group_lessons = GroupLesson.objects.filter(teacher=teacher).order_by('-date', '-start_time')

    context = {
        'group_lessons': group_lessons,
    }
    return render(request, 'school/teacher/group_lessons.html', context)


@login_required
def teacher_group_lesson_detail(request, lesson_id):
    """Детальная страница группового урока для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    lesson = get_object_or_404(GroupLesson, id=lesson_id)

    if lesson.teacher.user != request.user:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    enrollments = lesson.enrollments.all().select_related('student__user')

    context = {
        'lesson': lesson,
        'enrollments': enrollments,
    }
    return render(request, 'school/teacher/group_lesson_detail.html', context)


@login_required
@require_POST
def mark_group_attendance(request, lesson_id):
    """Отметить присутствие ученика на групповом уроке"""
    if request.user.role != 'teacher':
        return JsonResponse({'error': 'Доступ запрещен'}, status=403)

    lesson = get_object_or_404(GroupLesson, id=lesson_id)

    if lesson.teacher.user != request.user:
        return JsonResponse({'error': 'Доступ запрещен'}, status=403)

    enrollment_id = request.POST.get('enrollment_id')
    status = request.POST.get('status')

    enrollment = get_object_or_404(GroupEnrollment, id=enrollment_id, group_lesson=lesson)
    enrollment.status = status
    enrollment.save()

    return JsonResponse({'success': True})


@login_required
@require_POST
def complete_group_lesson(request, lesson_id):
    """Завершить групповой урок"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    lesson = get_object_or_404(GroupLesson, id=lesson_id)

    if lesson.teacher.user != request.user:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    if lesson.status != 'scheduled':
        messages.error(request, 'Урок уже завершен или отменен')
        return redirect('teacher_group_lesson_detail', lesson_id=lesson.id)

    lesson.mark_as_completed()

    messages.success(request, 'Групповой урок завершен')
    return redirect('teacher_group_lessons')


# ============================================
# ЧАСТЬ 11: ШАБЛОНЫ РАСПИСАНИЯ
# ============================================

@login_required
def teacher_schedule_templates(request):
    """Список шаблонов расписания учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    templates = ScheduleTemplate.objects.filter(teacher=teacher).order_by('-created_at')

    context = {
        'templates': templates,
    }
    return render(request, 'school/teacher/schedule_templates.html', context)


@login_required
def teacher_schedule_template_create(request):
    """Создание шаблона расписания"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    if request.method == 'POST':
        form = ScheduleTemplateForm(request.POST)
        if form.is_valid():
            template = form.save(commit=False)
            template.teacher = teacher
            template.save()
            form.save_m2m()

            messages.success(request, 'Шаблон расписания создан')
            return redirect('teacher_schedule_templates')
    else:
        form = ScheduleTemplateForm()
        form.fields['students'].queryset = teacher.student_set.all()

    context = {
        'form': form,
        'teacher': teacher,
    }
    return render(request, 'school/teacher/schedule_template_form.html', context)


@login_required
def teacher_schedule_template_detail(request, template_id):
    """Детали шаблона и генерация уроков"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    template = get_object_or_404(ScheduleTemplate, id=template_id, teacher=teacher)

    if request.method == 'POST' and 'generate' in request.POST:
        student_ids = request.POST.getlist('students')
        students = Student.objects.filter(id__in=student_ids, teachers=teacher)

        lessons = template.generate_lessons(students)
        messages.success(request, f'Создано {len(lessons)} уроков')
        return redirect('teacher_schedule_template_detail', template_id=template.id)

    context = {
        'template': template,
        'students': teacher.student_set.all(),
    }
    return render(request, 'school/teacher/schedule_template_detail.html', context)


@login_required
def teacher_schedule_template_delete(request, template_id):
    """Удаление шаблона расписания"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    template = get_object_or_404(ScheduleTemplate, id=template_id, teacher=teacher)

    if request.method == 'POST':
        template.delete()
        messages.success(request, 'Шаблон успешно удален')
        return redirect('teacher_schedule_templates')

    context = {
        'template': template,
    }
    return render(request, 'school/teacher/schedule_template_confirm_delete.html', context)


# ============================================
# ЧАСТЬ 12: ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================

@login_required
def teacher_edit_lesson(request, lesson_id):
    """Редактирование урока учителем"""
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if request.user.role != 'teacher' or lesson.teacher.user != request.user:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    if request.method == 'POST':
        lesson.date = request.POST.get('date')
        lesson.start_time = request.POST.get('start_time')
        lesson.end_time = request.POST.get('end_time')
        lesson.meeting_link = request.POST.get('meeting_link')
        lesson.meeting_platform = request.POST.get('meeting_platform')
        lesson.notes = request.POST.get('notes')
        lesson.save()

        messages.success(request, 'Урок обновлен')
        return redirect('teacher_lesson_detail', lesson_id=lesson.id)

    context = {
        'lesson': lesson,
    }
    return render(request, 'school/teacher/edit_lesson.html', context)

def check_teacher_busy(teacher, date, start_time, end_time, exclude_lesson_id=None):
    """
    Проверяет, занят ли учитель в указанное время
    Возвращает True если свободен, False если занят
    """
    from datetime import datetime

    # Проверяем существующие уроки
    existing_lessons = Lesson.objects.filter(
        teacher=teacher,
        date=date,
        status__in=['scheduled', 'completed']
    )

    if exclude_lesson_id:
        existing_lessons = existing_lessons.exclude(pk=exclude_lesson_id)

    lesson_start = datetime.combine(date, start_time)
    lesson_end = datetime.combine(date, end_time)

    for lesson in existing_lessons:
        existing_start = datetime.combine(lesson.date, lesson.start_time)
        existing_end = datetime.combine(lesson.date, lesson.end_time)

        if lesson_start < existing_end and lesson_end > existing_start:
            return False, lesson  # Занят, возвращаем конфликтующий урок

    return True, None  # Свободен


@login_required
def teacher_create_schedule(request):
    """Создание шаблона расписания (разового или повторяющегося)"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    if request.method == 'POST':
        student_id = request.POST.get('student')
        subject_id = request.POST.get('subject')
        topic = request.POST.get('topic', '')
        start_time_str = request.POST.get('start_time')
        end_time_str = request.POST.get('end_time')
        repeat_type = request.POST.get('repeat_type', 'single')
        notes = request.POST.get('notes', '')

        if not student_id or not subject_id or not start_time_str:
            messages.error(request, 'Заполните все обязательные поля')
            return redirect('teacher_create_schedule')

        student = get_object_or_404(Student, id=student_id, teachers=teacher)
        subject = get_object_or_404(Subject, id=subject_id)

        cost, teacher_payment = StudentSubjectPrice.get_price_for(student, subject, teacher)

        try:
            start_time = datetime.strptime(start_time_str, '%H:%M').time()
        except ValueError:
            messages.error(request, 'Неверный формат времени начала')
            return redirect('teacher_create_schedule')

        if not end_time_str:
            today_date = date.today()
            start_dt = datetime.combine(today_date, start_time)
            end_dt = start_dt + timedelta(hours=1)
            end_time = end_dt.time()
        else:
            try:
                end_time = datetime.strptime(end_time_str, '%H:%M').time()
            except ValueError:
                messages.error(request, 'Неверный формат времени окончания')
                return redirect('teacher_create_schedule')

        # ✅ ПРОВЕРКА НА ЗАНЯТОСТЬ УЧИТЕЛЯ
        if repeat_type == 'single':
            date_str = request.POST.get('date')
            if not date_str:
                messages.error(request, 'Укажите дату занятия')
                return redirect('teacher_create_schedule')

            lesson_date = datetime.strptime(date_str, '%Y-%m-%d').date()

            # Проверяем, свободен ли учитель
            is_free, conflict = check_teacher_busy(teacher, lesson_date, start_time, end_time)
            if not is_free:
                messages.error(
                    request,
                    f'❌ Учитель уже занят в это время!\n'
                    f'Конфликт с уроком: {conflict.subject.name}\n'
                    f'Время: {conflict.start_time} - {conflict.end_time}'
                )
                return redirect('teacher_create_schedule')
        else:
            # Для повторяющихся уроков проверяем первую дату (предварительно)
            start_date_str = request.POST.get('start_date')
            if start_date_str:
                first_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
                is_free, conflict = check_teacher_busy(teacher, first_date, start_time, end_time)
                if not is_free:
                    messages.warning(
                        request,
                        f'⚠️ Внимание! На дату {first_date.strftime("%d.%m.%Y")} '
                        f'учитель уже занят ({conflict.subject.name} в {conflict.start_time}).\n'
                        f'Остальные даты будут проверены при генерации.'
                    )

        template = ScheduleTemplate(
            teacher=teacher,
            subject=subject,
            start_time=start_time,
            end_time=end_time,
            repeat_type=repeat_type,
            notes=notes,
            base_cost=cost or Decimal('1000'),
            base_teacher_payment=teacher_payment or (cost or Decimal('1000')) * Decimal('0.7')
        )

        if repeat_type == 'single':
            date_str = request.POST.get('date')
            if not date_str:
                messages.error(request, 'Укажите дату занятия')
                return redirect('teacher_create_schedule')

            template.start_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            template.end_date = None
            template.max_occurrences = 1

        else:
            weekdays = request.POST.getlist('weekdays[]')
            start_date_str = request.POST.get('start_date')
            end_date_str = request.POST.get('end_date')
            max_occurrences = request.POST.get('max_occurrences')

            if not start_date_str:
                messages.error(request, 'Укажите дату начала расписания')
                return redirect('teacher_create_schedule')

            if not weekdays:
                messages.error(request, 'Выберите хотя бы один день недели')
                return redirect('teacher_create_schedule')

            template.start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            template.end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date() if end_date_str else None
            template.max_occurrences = int(max_occurrences) if max_occurrences else None

            template.monday = '1' in weekdays
            template.tuesday = '2' in weekdays
            template.wednesday = '3' in weekdays
            template.thursday = '4' in weekdays
            template.friday = '5' in weekdays
            template.saturday = '6' in weekdays
            template.sunday = '7' in weekdays

        template.save()
        template.students.add(student)

        # ✅ ГЕНЕРИРУЕМ УРОКИ
        lessons = template.generate_lessons()

        # ✅ ДОПОЛНИТЕЛЬНАЯ ПРОВЕРКА после генерации
        conflicts_after = []
        for lesson in lessons:
            is_free, conflict = check_teacher_busy(teacher, lesson.date, lesson.start_time, lesson.end_time, lesson.id)
            if not is_free:
                conflicts_after.append(lesson)
                lesson.delete()  # Удаляем конфликтующий урок

        if conflicts_after:
            messages.warning(
                request,
                f'⚠️ Обнаружены пересечения! Удалено {len(conflicts_after)} уроков.\n'
                f'Проверьте расписание учителя.'
            )

        # ✅ ОТПРАВЛЯЕМ УВЕДОМЛЕНИЯ ДЛЯ КАЖДОГО СОЗДАННОГО УРОКА
        if lessons:
            print(f"\n{'=' * 50}")
            print(f"📅 Создано уроков: {len(lessons)}")

            from school.telegram import notify_new_lesson
            from .models import Notification

            for lesson in lessons:
                # Внутренние уведомления
                print(f"\n✅ Обработка урока {lesson.id}")
                print(f"   Ученик: {student.user.get_full_name()}")

                # Уведомление ученику
                Notification.objects.create(
                    user=student.user,
                    title='📚 Новый урок',
                    message=f'Учитель {teacher.user.get_full_name()} назначил урок по {subject.name} на {lesson.date.strftime("%d.%m.%Y")} в {lesson.start_time.strftime("%H:%M")}',
                    notification_type='lesson_reminder',
                    link=f'/lesson/{lesson.id}/'
                )
                print(f"✅ Внутреннее уведомление ученику создано")

                # Уведомление учителю
                Notification.objects.create(
                    user=request.user,
                    title='✅ Урок создан',
                    message=f'Урок по {subject.name} для {student.user.get_full_name()} создан на {lesson.date.strftime("%d.%m.%Y")} в {lesson.start_time.strftime("%H:%M")}',
                    notification_type='system',
                    link=f'/teacher/lesson/{lesson.id}/'
                )
                print(f"✅ Внутреннее уведомление учителю создано")

                # Telegram уведомление
                try:
                    notify_new_lesson(lesson)
                    print(f"✅ Telegram уведомление отправлено для урока {lesson.id}")
                except Exception as e:
                    print(f"❌ Ошибка отправки Telegram: {e}")

            print(f"{'=' * 50}\n")

        if repeat_type == 'single':
            messages.success(request, f'✅ Урок создан на {template.start_date} в {start_time_str}')
        else:
            messages.success(request, f'✅ Расписание создано! Сгенерировано {len(lessons)} уроков')

        return redirect('teacher_dashboard')

    # GET запрос - показываем форму
    students = teacher.student_set.all()
    subjects = teacher.subjects.all()

    context = {
        'teacher': teacher,
        'students': students,
        'subjects': subjects,
        'today': timezone.now().date().strftime('%Y-%m-%d'),
    }
    return render(request, 'school/teacher/schedule_template_form.html', context)

# ✅ ВСПОМОГАТЕЛЬНАЯ ФУНКЦИЯ (добавьте в начало файла или перед функцией)
def check_teacher_busy(teacher, date, start_time, end_time, exclude_lesson_id=None):
    """
    Проверяет, занят ли учитель в указанное время
    Возвращает (True, None) если свободен, (False, conflict_lesson) если занят
    """
    from datetime import datetime

    existing_lessons = Lesson.objects.filter(
        teacher=teacher,
        date=date,
        status__in=['scheduled', 'completed']
    )

    if exclude_lesson_id:
        existing_lessons = existing_lessons.exclude(pk=exclude_lesson_id)

    lesson_start = datetime.combine(date, start_time)
    lesson_end = datetime.combine(date, end_time)

    for lesson in existing_lessons:
        existing_start = datetime.combine(lesson.date, lesson.start_time)
        existing_end = datetime.combine(lesson.date, lesson.end_time)

        # Проверка на пересечение интервалов
        if lesson_start < existing_end and lesson_end > existing_start:
            return False, lesson

    return True, None
@login_required
def profile(request):
    """Профиль пользователя"""
    if request.method == 'POST':
        form = ProfileUpdateForm(request.POST, request.FILES, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Профиль обновлен')
            return redirect('profile')
    else:
        form = ProfileUpdateForm(instance=request.user)

    return render(request, 'school/profile.html', {'form': form})


@login_required
def student_materials(request):
    """Все методические материалы для ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    student = request.user.student_profile
    teachers = student.teachers.all()

    # Материалы, доступные ученику:
    # 1. Публичные материалы
    # 2. Материалы, назначенные конкретно этому ученику
    # 3. Материалы учителей этого ученика
    materials = Material.objects.filter(
        Q(is_public=True) |
        Q(students=student) |
        Q(teachers__in=teachers)
    ).distinct().order_by('-created_at')

    # Фильтры
    subject_id = request.GET.get('subject')
    if subject_id:
        materials = materials.filter(subjects__id=subject_id)

    material_type = request.GET.get('type')
    if material_type:
        materials = materials.filter(material_type=material_type)

    teacher_id = request.GET.get('teacher')
    if teacher_id:
        materials = materials.filter(teachers__id=teacher_id)

    subjects = Subject.objects.all()

    context = {
        'materials': materials,
        'subjects': subjects,
        'teachers': teachers,
        'student': student,
    }
    return render(request, 'school/student/materials.html', context)

@login_required
def teacher_materials(request):
    """Управление методическими материалами для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    materials = Material.objects.filter(
        Q(teachers=teacher) | Q(created_by=request.user)
    ).distinct().order_by('-created_at')

    student_id = request.GET.get('student')
    if student_id:
        materials = materials.filter(students__id=student_id)

    students = teacher.student_set.all()

    context = {
        'materials': materials,
        'students': students,
        'teacher': teacher,
    }

    return render(request, 'school/teacher/materials.html', context)


@login_required
def teacher_student_detail(request, student_id):
    """Детальная информация об ученике для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    student = get_object_or_404(Student, id=student_id, teachers=teacher)

    # ИСПОЛЬЗУЕМ StudentFinanceHelper
    finance_helper = StudentFinanceHelper(student)

    lessons = Lesson.objects.filter(
        teacher=teacher,
        attendance__student=student
    ).select_related('subject', 'format').distinct().order_by('-date')

    notes = StudentNote.objects.filter(teacher=teacher, student=student).order_by('-created_at')

    materials = Material.objects.filter(
        Q(students=student) | Q(is_public=True)
    ).distinct()

    context = {
        'student': student,
        'finance': {
            'balance': float(finance_helper.balance),
            'stats': finance_helper.get_lessons_stats(30)
        },
        'lessons': lessons[:20],
        'notes': notes,
        'materials': materials,
    }

    return render(request, 'school/teacher/student_detail.html', context)


@login_required
def student_calendar(request):
    """Календарь ученика"""
    if request.user.role != 'student':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    student = request.user.student_profile

    lessons = Lesson.objects.filter(student=student).order_by('date', 'start_time')

    calendar_events = []
    for lesson in lessons:
        calendar_events.append({
            'title': f"{lesson.subject.name} - {lesson.teacher.user.get_full_name()}",
            'start': f"{lesson.date}T{lesson.start_time}",
            'end': f"{lesson.date}T{lesson.end_time}",
            'url': f"/lessons/{lesson.id}/",
            'status': lesson.status,
        })

    context = {
        'calendar_events': calendar_events,
    }
    return render(request, 'school/student/calendar.html', context)


# ============================================
# ЧАСТЬ 13: ПОДТВЕРЖДЕНИЕ EMAIL
# ============================================

def verify_email(request, token):
    """Подтверждение email по токену"""
    print(f"\n{'=' * 50}")
    print(f"🔍 verify_email вызван с токеном: {token}")
    print(f"{'=' * 50}\n")

    try:
        verification_token = get_object_or_404(EmailVerificationToken, token=token)

        if not verification_token.is_valid():
            messages.error(
                request,
                'Срок действия ссылки истек. Запросите повторную отправку письма.'
            )
            return redirect('resend_verification')

        user = verification_token.user

        if user.is_email_verified:
            messages.info(request, 'Email уже подтвержден')
            return redirect('login')

        user.is_email_verified = True
        user.save(update_fields=['is_email_verified'])

        try:
            send_verification_success_email(user)
        except Exception as e:
            logger.error(f"Ошибка отправки письма об успехе: {e}")

        verification_token.delete()

        messages.success(
            request,
            '✅ Email успешно подтвержден! Теперь вы можете войти в систему.'
        )

    except EmailVerificationToken.DoesNotExist:
        messages.error(request, '❌ Недействительная ссылка подтверждения')
    except Exception as e:
        traceback.print_exc()
        messages.error(request, f'❌ Ошибка при подтверждении: {str(e)}')

    return redirect('login')


def resend_verification(request):
    """Повторная отправка письма подтверждения"""
    if request.method == 'POST':
        email = request.POST.get('email')

        try:
            user = User.objects.get(email=email)

            if user.is_email_verified:
                messages.info(
                    request,
                    'Этот email уже подтвержден. Вы можете войти в систему.'
                )
                return redirect('login')

            if user.email_verification_sent:
                time_since = timezone.now() - user.email_verification_sent
                if time_since.total_seconds() < 300:
                    minutes_left = 5 - (time_since.total_seconds() // 60)
                    messages.error(
                        request,
                        f'Письмо уже отправлено. Повторная отправка через {int(minutes_left)} минут'
                    )
                    return redirect('login')

            if send_verification_email(user, request):
                messages.success(
                    request,
                    'Письмо с подтверждением отправлено повторно. Проверьте вашу почту.'
                )
            else:
                messages.error(
                    request,
                    'Ошибка при отправке письма. Попробуйте позже.'
                )

        except User.DoesNotExist:
            messages.success(
                request,
                'Если пользователь с таким email существует, письмо будет отправлено повторно.'
            )

    return render(request, 'school/resend_verification.html')


@login_required
@require_POST
def complete_lesson(request, lesson_id):
    """Завершение урока и создание отчета с учетом явки"""
    lesson = get_object_or_404(Lesson, id=lesson_id)

    if request.user.role != 'teacher' or lesson.teacher.user != request.user:
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    if lesson.status != 'scheduled':
        messages.error(request, 'Урок уже завершен или отменен')
        return redirect('teacher_lesson_detail', lesson_id=lesson.id)

    report_data = {
        'topic': request.POST.get('topic'),
        'covered_material': request.POST.get('covered_material'),
        'homework': request.POST.get('homework'),
        'student_progress': request.POST.get('student_progress'),
        'next_lesson_plan': request.POST.get('next_lesson_plan', '')
    }

    required_fields = ['topic', 'covered_material', 'homework', 'student_progress']
    if not all([report_data.get(field) for field in required_fields]):
        messages.error(request, 'Заполните все обязательные поля')
        return redirect('teacher_lesson_detail', lesson_id=lesson.id)

    attended_students = []
    for attendance in lesson.attendance.all():
        if request.POST.get(f'attended_{attendance.id}'):
            attended_students.append(attendance.id)
            attendance.status = 'attended'
            attendance.save()
        else:
            attendance.status = 'absent'
            attendance.save()

    # ✅ Создание отчета (триггерит сигнал lesson_completed_notifications)
    report = lesson.mark_as_completed(report_data, attended_students)

    if report:
        messages.success(request,
                         f'Урок завершен. Отчет #{report.id} создан. Присутствовало: {len(attended_students)} учеников.')
    else:
        messages.success(request, f'Урок завершен. Присутствовало: {len(attended_students)} учеников.')

    return redirect('teacher_lesson_detail', lesson_id=lesson.id)


@login_required
def telegram_settings(request):
    """Настройки Telegram уведомлений"""
    if request.method == 'POST':
        form = TelegramSettingsForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Настройки Telegram сохранены!')
            return redirect('profile')
    else:
        form = TelegramSettingsForm(instance=request.user)

    # Инструкция как получить Chat ID
    bot_username = 'PlusProgressBot'  # имя вашего бота
    instruction = f"""
    1. Напишите боту @{bot_username} любое сообщение
    2. Перейдите по ссылке: https://api.telegram.org/bot{settings.TELEGRAM_BOT_TOKEN}/getUpdates
    3. Найдите в ответе ваш "id" (число) и скопируйте его
    """

    context = {
        'form': form,
        'instruction': instruction,
        'bot_username': bot_username,
    }
    return render(request, 'school/telegram_settings.html', context)



User = get_user_model()


@csrf_exempt
def telegram_webhook(request):
    """Обработчик входящих сообщений от Telegram"""
    print(f"\n🔥🔥🔥 ВЫЗВАНА ФУНКЦИЯ WEBHOOK 🔥🔥🔥")
    print(f"Метод запроса: {request.method}")
    print(f"Тело запроса: {request.body}")

    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            print(f"\n{'=' * 50}")
            print(f"📱 ПОЛУЧЕНО СООБЩЕНИЕ ОТ TELEGRAM")
            print(json.dumps(data, indent=2, ensure_ascii=False))
            print(f"{'=' * 50}\n")

            # Получаем информацию о пользователе
            if 'message' in data:
                chat_id = data['message']['chat']['id']
                username = data['message']['from'].get('username', '')
                first_name = data['message']['from'].get('first_name', '')
                last_name = data['message']['from'].get('last_name', '')
                text = data['message'].get('text', '')

                print(f"✅ Chat ID: {chat_id}")
                print(f"✅ Username: @{username}")
                print(f"✅ Имя: {first_name} {last_name}")
                print(f"✅ Текст: {text}")

                # ✅ Отправляем ответ через Telegram API
                bot_token = settings.TELEGRAM_BOT_TOKEN
                url = f"https://api.telegram.org/bot{bot_token}/sendMessage"

                response_text = f"Привет, {first_name}! Твой Telegram ID: {chat_id}\n\n"
                response_text += "Скопируй этот ID и вставь в настройках профиля на сайте, чтобы получать уведомления."

                response = requests.post(url, json={
                    'chat_id': chat_id,
                    'text': response_text,
                    'parse_mode': 'HTML'
                })

                print(f"✅ Ответ отправлен, статус: {response.status_code}")
                print(f"✅ Ответ API: {response.text}")

                return JsonResponse({'ok': True})

        except Exception as e:
            print(f"❌ Ошибка: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'ok': False, 'error': str(e)})

    return JsonResponse({'ok': True})


# school/views.py

@login_required
def teacher_student_detail(request, student_id):
    """Детальная информация об ученике для учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile
    student = get_object_or_404(Student, id=student_id, teachers=teacher)

    # Получаем уроки ученика с этим учителем
    lessons = Lesson.objects.filter(
        teacher=teacher,
        attendance__student=student
    ).select_related('subject', 'format').distinct().order_by('-date')

    # Получаем заметки об ученике
    notes = StudentNote.objects.filter(teacher=teacher, student=student).order_by('-created_at')

    # Получаем домашние задания
    homeworks = Homework.objects.filter(student=student, teacher=teacher).order_by('-created_at')

    # Статистика
    total_lessons = lessons.count()
    completed_lessons = lessons.filter(status='completed').count()

    context = {
        'student': student,
        'lessons': lessons[:20],
        'notes': notes,
        'homeworks': homeworks[:10],
        'total_lessons': total_lessons,
        'completed_lessons': completed_lessons,
    }
    return render(request, 'school/teacher/student_detail.html', context)


@login_required
def teacher_payments(request):
    """Страница со всеми выплатами учителя"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    # Все выплаты учителя
    payments = Payment.objects.filter(
        user=request.user,
        payment_type='teacher_payment'
    ).order_by('-created_at')

    # Статистика по годам
    import calendar
    from django.db.models import Sum
    from datetime import datetime

    current_year = datetime.now().year
    years = range(current_year - 2, current_year + 1)

    monthly_stats = []
    for year in years:
        for month in range(1, 13):
            month_payments = payments.filter(
                created_at__year=year,
                created_at__month=month
            )
            total = month_payments.aggregate(Sum('amount'))['amount__sum'] or 0
            count = month_payments.count()

            if total > 0 or count > 0:
                monthly_stats.append({
                    'year': year,
                    'month': month,
                    'month_name': calendar.month_name[month],
                    'total': total,
                    'count': count,
                })

    context = {
        'payments': payments,
        'monthly_stats': monthly_stats,
        'total_payments': payments.aggregate(Sum('amount'))['amount__sum'] or 0,
    }
    return render(request, 'school/teacher/payments.html', context)


@login_required
def teacher_request_payment(request):
    """Запрос выплаты учителем"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    if request.method == 'POST':
        amount = request.POST.get('amount')
        payment_method = request.POST.get('payment_method')
        payment_details = request.POST.get('payment_details')

        try:
            amount = Decimal(amount)
            if amount <= 0:
                messages.error(request, 'Сумма должна быть положительной')
                return redirect('teacher_request_payment')

            if amount > teacher.wallet_balance:
                messages.error(request, 'Недостаточно средств на балансе')
                return redirect('teacher_request_payment')

            # ✅ СОЗДАЕМ ЗАПРОС В БАЗЕ ДАННЫХ
            payment_request = PaymentRequest.objects.create(
                teacher=teacher,
                amount=amount,
                payment_method=payment_method,
                payment_details=payment_details,
                status='pending'
            )

            # Уведомление админу
            admin_users = User.objects.filter(role='admin')
            for admin in admin_users:
                Notification.objects.create(
                    user=admin,
                    title='💰 Новый запрос выплаты',
                    message=f'Учитель {teacher.user.get_full_name()} запросил выплату {amount} ₽',
                    notification_type='system',
                    link=f'/admin/school/paymentrequest/{payment_request.id}/change/'
                )

            # Уведомление учителю
            Notification.objects.create(
                user=request.user,
                title='✅ Запрос отправлен',
                message=f'Запрос на выплату {amount} ₽ отправлен администратору. Номер запроса: #{payment_request.id}',
                notification_type='payment_withdrawn',
                link='/teacher/dashboard/#payments'
            )

            # Telegram уведомление
            try:
                from school.telegram import send_telegram_message
                admin_text = (
                    f"💰 НОВЫЙ ЗАПРОС ВЫПЛАТЫ #{payment_request.id}\n\n"
                    f"Учитель: {teacher.user.get_full_name()}\n"
                    f"Сумма: {amount} ₽\n"
                    f"Способ: {payment_method}\n"
                    f"Ссылка: /admin/school/paymentrequest/{payment_request.id}/change/"
                )
                send_telegram_message(admin_text)
            except:
                pass

            messages.success(request, f'✅ Запрос #{payment_request.id} на выплату {amount} ₽ отправлен администратору')
            return redirect('teacher_dashboard')

        except Exception as e:
            messages.error(request, f'Ошибка: {str(e)}')
            return redirect('teacher_request_payment')

    # GET запрос - показываем форму
    real_payments = Payment.objects.filter(
        user=request.user,
        payment_type='teacher_payment'
    ).order_by('-created_at')[:5]

    # Показываем последние запросы учителя
    recent_requests = PaymentRequest.objects.filter(
        teacher=teacher
    ).order_by('-created_at')[:5]

    from .models import Lesson
    lessons = Lesson.objects.filter(teacher=teacher, status='completed')
    calculator = PeriodFinanceCalculator(lessons)

    context = {
        'teacher': teacher,
        'real_payments': real_payments,
        'recent_requests': recent_requests,
        'stats': {
            'total_earned': calculator.lessons_stats['teacher_payment'],
            'paid': calculator.payments_stats['teacher_payments'],
            'available': teacher.wallet_balance,
        }
    }
    return render(request, 'school/teacher/payment_request.html', context)


@login_required
@login_required
def teacher_material_create(request):
    """Создание методического материала (для учеников)"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    teacher = request.user.teacher_profile

    if request.method == 'POST':
        title = request.POST.get('title')
        description = request.POST.get('description')
        material_type = request.POST.get('material_type')
        file = request.FILES.get('file')
        link = request.POST.get('link')
        subjects = request.POST.getlist('subjects')
        students = request.POST.getlist('students')  # список ID учеников
        is_public = request.POST.get('is_public') == 'on'

        if not title or not material_type:
            messages.error(request, 'Заполните обязательные поля')
            return redirect('teacher_material_create')

        if material_type == 'file' and not file:
            messages.error(request, 'Выберите файл для загрузки')
            return redirect('teacher_material_create')

        if material_type == 'link' and not link:
            messages.error(request, 'Введите ссылку')
            return redirect('teacher_material_create')

        # Создаем материал
        material = Material.objects.create(
            title=title,
            description=description,
            material_type=material_type,
            file=file,
            link=link,
            created_by=request.user,
            is_public=is_public
        )

        # Добавляем предметы
        if subjects:
            material.subjects.set(subjects)

        # Добавляем учителя как владельца
        material.teachers.add(teacher)

        # Добавляем выбранных учеников
        if students and not is_public:
            material.students.set(students)

        # Логирование
        log_user_action(
            request,
            'material_add',
            f'Добавлен методический материал: {title}',
            object_id=material.id,
            object_type='material'
        )

        # Уведомления ученикам
        if students:
            for student_id in students:
                try:
                    student = Student.objects.get(id=student_id)
                    Notification.objects.create(
                        user=student.user,
                        title='📚 Новый учебный материал',
                        message=f'Учитель {teacher.user.get_full_name()} добавил материал: {title}',
                        notification_type='material_added',
                        link=f'/student/materials/'
                    )
                except Student.DoesNotExist:
                    pass

        messages.success(request, f'✅ Материал "{title}" успешно добавлен')
        return redirect('teacher_dashboard')

    # GET запрос - показываем форму
    students = teacher.student_set.all().select_related('user')
    subjects = Subject.objects.all()

    context = {
        'teacher': teacher,
        'students': students,
        'subjects': subjects,
    }
    return render(request, 'school/teacher/material_form.html', context)


@login_required
def teacher_material_edit(request, material_id):
    """Редактирование методического материала"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    material = get_object_or_404(Material, id=material_id, created_by=request.user)
    teacher = request.user.teacher_profile

    if request.method == 'POST':
        material.title = request.POST.get('title')
        material.description = request.POST.get('description')
        material.is_public = request.POST.get('is_public') == 'on'

        subjects = request.POST.getlist('subjects')
        students = request.POST.getlist('students')

        if request.FILES.get('file'):
            material.file = request.FILES['file']

        if request.POST.get('link'):
            material.link = request.POST.get('link')

        material.save()

        # Обновляем связи
        if subjects:
            material.subjects.set(subjects)

        # Обновляем учеников
        if not material.is_public and students:
            material.students.set(students)
        elif material.is_public:
            material.students.clear()  # для публичных материалов ученики не нужны

        messages.success(request, f'✅ Материал "{material.title}" обновлен')
        return redirect('teacher_dashboard')

    students = teacher.student_set.all().select_related('user')
    subjects = Subject.objects.all()

    context = {
        'material': material,
        'students': students,
        'subjects': subjects,
    }
    return render(request, 'school/teacher/material_form.html', context)
@login_required
def teacher_material_delete(request, material_id):
    """Удаление методического материала"""
    if request.user.role != 'teacher':
        messages.error(request, 'Доступ запрещен')
        return redirect('dashboard')

    material = get_object_or_404(Material, id=material_id, created_by=request.user)

    if request.method == 'POST':
        title = material.title
        material.delete()
        messages.success(request, f'✅ Материал "{title}" удален')
        return redirect('teacher_dashboard')

    context = {'material': material}
    return render(request, 'school/teacher/material_confirm_delete.html', context)


@login_required
def material_detail(request, material_id):
    """Просмотр материала"""
    material = get_object_or_404(Material, id=material_id)

    # Проверка доступа
    if not material.is_public:
        if request.user.role == 'student':
            student = request.user.student_profile
            if student not in material.students.all():
                messages.error(request, 'Доступ запрещен')
                return redirect('dashboard')
        elif request.user.role == 'teacher':
            teacher = request.user.teacher_profile
            if teacher not in material.teachers.all() and material.created_by != request.user:
                messages.error(request, 'Доступ запрещен')
                return redirect('dashboard')

    # Увеличиваем счетчик просмотров (можно добавить поле views)
    # material.views += 1
    # material.save()

    context = {
        'material': material,
    }
    return render(request, 'school/material_detail.html', context)




'''Отдельные представления страниц сайта'''


def article_artikli(request):
    return render(request, 'school/articles/artikli-v-anglijskom-yazyke.html')


def article_glagol(request):
    return render(request, 'school/articles/glagol-to-be-v-anglijskom-yazyke.html')


def article_kolichestvennye(request):
    return render(request, 'school/articles/kolichestvennye-mestoimeniya-v-anglijskom-yazyke.html')


def article_mestoimeniya(request):
    return render(request, 'school/articles/mestoimeniya-v-anglijskom-yazyke.html')




@require_POST
def trial_request(request):
    """Обработка заявки на пробный урок (старая версия, скоро будет удалена)"""
    """ODO: Удалить после проверки trial_request_ajax"""
    
    # Проверяем, AJAX ли это запрос
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # AJAX запрос - возвращаем JSON
        name = request.POST.get('name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        subject = request.POST.get('subject')
        
        TrialRequest.objects.create(
            name=name,
            email=email,
            phone=phone,
            subject=subject
        )
        
        return JsonResponse({'status': 'ok', 'message': 'Заявка отправлена!'})
    
    else:
        # Обычный POST запрос - редирект (для обратной совместимости)
        name = request.POST.get('name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        subject = request.POST.get('subject')
        
        TrialRequest.objects.create(
            name=name,
            email=email,
            phone=phone,
            subject=subject
        )
        
        messages.success(request, 'Спасибо! Ваша заявка принята. Мы свяжемся с вами в ближайшее время.')
        return redirect(request.META.get('HTTP_REFERER', 'home'))
    """Обработка заявки на пробный урок"""
    name = request.POST.get('name')
    email = request.POST.get('email')
    phone = request.POST.get('phone')
    subject = request.POST.get('subject')

    TrialRequest.objects.create(
        name=name,
        email=email,
        phone=phone,
        subject=subject
    )

    messages.success(request, 'Спасибо! Ваша заявка принята. Мы свяжемся с вами в ближайшее время.')
    return redirect(request.META.get('HTTP_REFERER', 'home'))



logger = logging.getLogger(__name__)

@require_POST
def trial_request_ajax(request):
    """
    AJAX обработка заявки на пробный урок
    """
    logger.info("="*50)
    logger.info("🔥 ПОЛУЧЕН AJAX ЗАПРОС НА ЗАЯВКУ")
    logger.info(f"POST данные: {request.POST}")
    try:
        # Получаем данные
        name = request.POST.get('name', '').strip()
        email = request.POST.get('email', '').strip()
        phone = request.POST.get('phone', '').strip()
        subject = request.POST.get('subject', '').strip()
        
        # Валидация
        if not name:
            return JsonResponse({'error': 'Укажите имя'}, status=400)
        if not phone:
            return JsonResponse({'error': 'Укажите телефон'}, status=400)
        if not subject:
            return JsonResponse({'error': 'Выберите предмет'}, status=400)
        
        # Сохраняем в базу
        trial = TrialRequest.objects.create(
            name=name,
            email=email,
            phone=phone,
            subject=subject
        )
        
        logger.info(f"Новая заявка #{trial.id} от {name}")
        
        # Отправляем email (опционально)
        try:
            from django.core.mail import send_mail
            send_mail(
                f'🔔 Новая заявка от {name}',
                f'Имя: {name}\nEmail: {email}\nТелефон: {phone}\nПредмет: {subject}',
                'jserge@yandex.ru',
                ['jserge@yandex.ru'],
                fail_silently=True,
            )
        except Exception as e:
            logger.error(f"Ошибка отправки email: {e}")
        
        return JsonResponse({'status': 'ok'})
        
    except Exception as e:
        logger.error(f"Ошибка в заявке: {e}")
        return JsonResponse({'error': 'Ошибка сервера'}, status=500)

# ============================================
# REST API VIEWSETS 
# ============================================
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.db.models import Q
from rest_framework import viewsets, permissions, status, generics
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.views import APIView
from django.contrib.auth import authenticate
from rest_framework.authtoken.models import Token
from .serializers import (
    TeacherSerializer, StudentSerializer, LessonSerializer, 
    SimpleLessonSerializer, UserSerializer, 
    HomeworkSerializer, HomeworkSubmissionSerializer,
    MaterialSerializer, PaymentSerializer,
    NotificationSerializer, LessonFeedbackSerializer,
    GroupLessonSerializer, GroupEnrollmentSerializer,
    ScheduleTemplateSerializer, StudentSubjectPriceSerializer,
    TrialRequestSerializer, DepositSerializer,
    StudentNoteSerializer, PaymentRequestSerializer, LessonReportSerializer
)


class IsTeacherUser(permissions.BasePermission):
    """Права для учителей"""
    def has_permission(self, request, view):
        return request.user.is_authenticated and request.user.role == 'teacher'

class IsStudentUser(permissions.BasePermission):
    """Права для учеников"""
    def has_permission(self, request, view):
        return request.user.is_authenticated and request.user.role == 'student'

class RegisterView(generics.CreateAPIView):
    """Регистрация нового пользователя"""
    serializer_class = RegisterSerializer
    permission_classes = [permissions.AllowAny]  # Разрешаем всем
    authentication_classes = []  # Можно вообще без аутентификации
    
    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            
            # Создаем профиль ученика автоматически
            Student.objects.create(user=user)
            
            # Создаем токен для пользователя
            from rest_framework.authtoken.models import Token
            token, created = Token.objects.get_or_create(user=user)
            
            return Response({
                'user': UserSerializer(user).data,
                'token': token.key,
                'message': 'Пользователь успешно зарегистрирован'
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class TeacherViewSet(viewsets.ModelViewSet):
    """API для учителей"""
    queryset = Teacher.objects.all().select_related('user').prefetch_related('subjects')
    serializer_class = TeacherSerializer
    permission_classes = [permissions.IsAdminUser]


class StudentViewSet(viewsets.ModelViewSet):
    """API для учеников"""
    queryset = Student.objects.all().select_related('user').prefetch_related('teachers')
    serializer_class = StudentSerializer
    permission_classes = [permissions.IsAdminUser]  # Для админов - полный доступ

    def get_queryset(self):
        queryset = super().get_queryset()
        user_id = self.request.query_params.get('user_id')
        if user_id:
            queryset = queryset.filter(user_id=user_id)
        return queryset
    
    @action(detail=False, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def me(self, request):
        """Получить данные текущего ученика"""
        try:
            student = Student.objects.get(user=request.user)
            serializer = self.get_serializer(student)
            return Response(serializer.data)
        except Student.DoesNotExist:
            return Response(
                {'detail': 'Профиль ученика не найден'}, 
                status=status.HTTP_404_NOT_FOUND
            )


class LessonViewSet(viewsets.ModelViewSet):
    """API для уроков"""
    queryset = Lesson.objects.all().select_related('teacher__user', 'subject')
    serializer_class = LessonSerializer

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [permissions.IsAdminUser | IsTeacherUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]

    def get_serializer_class(self):
        if self.action == 'list':
            return SimpleLessonSerializer
        return LessonSerializer

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        if user.role == 'teacher':
            queryset = queryset.filter(teacher__user=user)
        elif user.role == 'student':
            queryset = queryset.filter(attendance__student__user=user).distinct()

        # ... остальные фильтры ...
        return queryset

    def retrieve(self, request, *args, **kwargs):
        """Получение конкретного урока с проверкой прав"""
        lesson = self.get_object()

        # Проверка прав для ученика
        if request.user.role == 'student':
            if not lesson.attendance.filter(student__user=request.user).exists():
                return Response(
                    {"detail": "У вас нет доступа к этому уроку."},
                    status=status.HTTP_403_FORBIDDEN
                )

        # Проверка прав для учителя
        elif request.user.role == 'teacher':
            if lesson.teacher.user != request.user:
                return Response(
                    {"detail": "У вас нет доступа к этому уроку."},
                    status=status.HTTP_403_FORBIDDEN
                )

        serializer = self.get_serializer(lesson)
        return Response(serializer.data)

class LoginAPIView(APIView):
    permission_classes = [permissions.AllowAny]
    
    def post(self, request):
        import json
        print("="*50)
        print("🔥 LoginAPIView вызван")
        print(f"Content-Type: {request.content_type}")
        print(f"Данные запроса: {request.data}")
        print("="*50)
        
        username = request.data.get('username')
        password = request.data.get('password')
        
        if not username or not password:
            return Response(
                {'error': 'Не указаны username или password'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Аутентификация
        user = authenticate(username=username, password=password)
        
        if user:
            # Получаем или создаем токен
            token, created = Token.objects.get_or_create(user=user)
            return Response({
                'token': token.key,
                'user_id': user.id,
                'username': user.username,
                'role': user.role
            })
        
        return Response(
            {'error': 'Неверные учетные данные'}, 
            status=status.HTTP_401_UNAUTHORIZED
        )

class HomeworkViewSet(viewsets.ModelViewSet):
    """API для домашних заданий"""
    queryset = Homework.objects.all().select_related(
        'teacher__user', 'student__user', 'subject', 'lesson'
    )
    serializer_class = HomeworkSerializer
    
    def get_permissions(self):
        """Разные права для разных действий"""
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            permission_classes = [permissions.IsAdminUser | IsTeacherUser]
        else:
            permission_classes = [permissions.IsAuthenticated]
        return [permission() for permission in permission_classes]
    
    def get_queryset(self):
        """Пользователи видят только свои ДЗ"""
        user = self.request.user
        if user.is_anonymous:
            return Homework.objects.none()
        
        queryset = super().get_queryset()
        if user.role == 'student':
            queryset = queryset.filter(student__user=user)
        elif user.role == 'teacher':
            queryset = queryset.filter(teacher__user=user)
        
        # Фильтр по статусу (упрощенный вариант)
        status = self.request.query_params.get('status')
        if status:
            # Здесь нужна дополнительная логика фильтрации по статусу
            pass
        
        return queryset
    
    @action(detail=True, methods=['post'])
    def submit(self, request, pk=None):
        """Сдать домашнее задание"""
        homework = self.get_object()
        # Проверка, что задание принадлежит ученику
        if homework.student.user != request.user:
            return Response(
                {'error': 'Это не ваше задание'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Здесь логика сдачи ДЗ
        serializer = HomeworkSubmissionSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(homework=homework, student=homework.student)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class HomeworkSubmissionViewSet(viewsets.ModelViewSet):
    """API для сданных заданий"""
    queryset = HomeworkSubmission.objects.all().select_related(
        'homework', 'student__user'
    )
    serializer_class = HomeworkSubmissionSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        """Учителя видят все, ученики только свои"""
        user = self.request.user
        if user.role == 'student':
            return self.queryset.filter(student__user=user)
        elif user.role == 'teacher':
            return self.queryset.filter(homework__teacher__user=user)
        return self.queryset
    
    @action(detail=True, methods=['post'])
    def check(self, request, pk=None):
        """Проверить задание (для учителя)"""
        submission = self.get_object()
        if request.user.role != 'teacher':
            return Response(
                {'error': 'Только учитель может проверять'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        submission.grade = request.data.get('grade')
        submission.teacher_comment = request.data.get('teacher_comment')
        submission.status = 'checked'
        submission.checked_at = timezone.now()
        submission.save()
        
        return Response(HomeworkSubmissionSerializer(submission).data)

class MaterialViewSet(viewsets.ModelViewSet):
    """API для методических материалов"""
    queryset = Material.objects.all().prefetch_related('teachers', 'students', 'subjects')
    serializer_class = MaterialSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        """Фильтрация материалов по доступности"""
        user = self.request.user
        queryset = super().get_queryset()
        
        if user.role == 'student':
            # Ученики видят публичные материалы и свои
            student = user.student_profile
            return queryset.filter(
                Q(is_public=True) | 
                Q(students=student) |
                Q(teachers__in=student.teachers.all())
            ).distinct()
        elif user.role == 'teacher':
            # Учителя видят свои материалы
            teacher = user.teacher_profile
            return queryset.filter(
                Q(teachers=teacher) | 
                Q(created_by=user)
            ).distinct()
        
        return queryset

class PaymentViewSet(viewsets.ModelViewSet):
    """API для платежей"""
    queryset = Payment.objects.all().select_related('user', 'lesson')
    serializer_class = PaymentSerializer
    permission_classes = [permissions.IsAdminUser]
    
    def get_queryset(self):
        """Фильтрация по пользователю и типу"""
        queryset = super().get_queryset()
        user_id = self.request.query_params.get('user')
        payment_type = self.request.query_params.get('type')
        
        if user_id:
            queryset = queryset.filter(user_id=user_id)
        if payment_type:
            queryset = queryset.filter(payment_type=payment_type)
        
        return queryset

class NotificationViewSet(viewsets.ModelViewSet):
    """API для уведомлений"""
    serializer_class = NotificationSerializer
    permission_classes = [permissions.IsAuthenticated]
    queryset = Notification.objects.all()
    
    def get_queryset(self):
        """Пользователи видят только свои уведомления"""
        return Notification.objects.filter(user=self.request.user)
    
    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        notification = self.get_object()
        notification.is_read = True
        notification.save()
        return Response({'status': 'ok'})
    
    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        self.get_queryset().update(is_read=True)
        return Response({'status': 'ok'})

class LessonFeedbackViewSet(viewsets.ModelViewSet):
    """API для отзывов о уроках"""
    queryset = LessonFeedback.objects.all().select_related(
        'lesson', 'student__user', 'teacher__user'
    )
    serializer_class = LessonFeedbackSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        """Фильтрация по учителю/ученику"""
        user = self.request.user
        teacher_id = self.request.query_params.get('teacher')
        student_id = self.request.query_params.get('student')
        
        queryset = super().get_queryset()
        
        if teacher_id:
            queryset = queryset.filter(teacher_id=teacher_id)
        if student_id:
            queryset = queryset.filter(student_id=student_id)
        
        return queryset
    
    def perform_create(self, serializer):
        """Автоматически подставляем ученика"""
        lesson = serializer.validated_data['lesson']
        student = lesson.attendance.get(student__user=self.request.user).student
        serializer.save(student=student)

class GroupLessonViewSet(viewsets.ModelViewSet):
    """API для групповых уроков"""
    queryset = GroupLesson.objects.all().select_related('teacher__user', 'subject', 'format')
    serializer_class = GroupLessonSerializer
    permission_classes = [permissions.IsAdminUser]
    
    @action(detail=True, methods=['post'])
    def enroll(self, request, pk=None):
        """Записать ученика на групповой урок"""
        lesson = self.get_object()
        student_id = request.data.get('student_id')
        
        try:
            student = Student.objects.get(id=student_id)
            enrollment, created = GroupEnrollment.objects.get_or_create(
                group_lesson=lesson,
                student=student,
                defaults={'status': 'registered'}
            )
            return Response(GroupEnrollmentSerializer(enrollment).data)
        except Student.DoesNotExist:
            return Response(
                {'error': 'Ученик не найден'},
                status=status.HTTP_404_NOT_FOUND
            )

class GroupEnrollmentViewSet(viewsets.ModelViewSet):
    """API для записей на групповые уроки"""
    queryset = GroupEnrollment.objects.all().select_related(
        'group_lesson', 'student__user'
    )
    serializer_class = GroupEnrollmentSerializer
    permission_classes = [permissions.IsAdminUser]

class ScheduleTemplateViewSet(viewsets.ModelViewSet):
    """API для шаблонов расписания"""
    queryset = ScheduleTemplate.objects.all().select_related('teacher__user', 'subject', 'format')
    serializer_class = ScheduleTemplateSerializer
    permission_classes = [permissions.IsAdminUser]
    
    @action(detail=True, methods=['post'])
    def generate(self, request, pk=None):
        """Сгенерировать уроки из шаблона"""
        template = self.get_object()
        lessons = template.generate_lessons()
        return Response({
            'message': f'Создано {len(lessons)} уроков',
            'lessons': LessonSerializer(lessons, many=True).data
        })

class StudentSubjectPriceViewSet(viewsets.ModelViewSet):
    """API для индивидуальных цен"""
    queryset = StudentSubjectPrice.objects.all().select_related(
        'student__user', 'teacher__user', 'subject'
    )
    serializer_class = StudentSubjectPriceSerializer
    permission_classes = [permissions.IsAdminUser]

class TrialRequestViewSet(viewsets.ModelViewSet):
    """API для заявок на пробный урок"""
    queryset = TrialRequest.objects.all()
    serializer_class = TrialRequestSerializer
    
    def get_permissions(self):
        """На создание заявки могут все, остальное только админ"""
        if self.action == 'create':
            return [permissions.AllowAny()]
        return [permissions.IsAdminUser()]

class DepositViewSet(viewsets.ModelViewSet):
    """API для депозитов"""
    queryset = Deposit.objects.all().select_related('student__user', 'created_by')
    serializer_class = DepositSerializer
    permission_classes = [permissions.IsAdminUser]

class StudentNoteViewSet(viewsets.ModelViewSet):
    """API для заметок об учениках"""
    queryset = StudentNote.objects.all().select_related(
        'teacher__user', 'student__user'
    )
    serializer_class = StudentNoteSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        """Учителя видят свои заметки, ученики - заметки о себе"""
        user = self.request.user
        if user.role == 'teacher':
            return self.queryset.filter(teacher__user=user)
        elif user.role == 'student':
            return self.queryset.filter(student__user=user)
        return self.queryset

class PaymentRequestViewSet(viewsets.ModelViewSet):
    """API для запросов на выплаты"""
    queryset = PaymentRequest.objects.all().select_related('teacher__user', 'payment')
    serializer_class = PaymentRequestSerializer
    permission_classes = [permissions.IsAdminUser]

class LessonReportViewSet(viewsets.ModelViewSet):
    """API для отчетов о уроках"""
    queryset = LessonReport.objects.all().select_related('lesson')
    serializer_class = LessonReportSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        """Доступ к отчетам есть у учителя и учеников урока"""
        user = self.request.user
        queryset = super().get_queryset()
        
        if user.role == 'teacher':
            return queryset.filter(lesson__teacher__user=user)
        elif user.role == 'student':
            return queryset.filter(lesson__attendance__student__user=user)
        
        return queryset


def teachers_team(request):
    """
    Страница со всеми учителями школы
    """
    # Получаем только учителей с show_on_team=True
    teachers = Teacher.objects.filter(
        show_on_team=True
    ).select_related(
        'user'
    ).prefetch_related(
        'subjects'
    ).order_by('team_sort_order', '-team_highlight', 'user__last_name')

    # Разделяем на выделенных и обычных
    highlighted_teachers = []
    regular_teachers = []

    for teacher in teachers:
        # БЕЗОПАСНО получаем рейтинг
        try:
            rating_stats = teacher.rating_stats
        except:
            # Если рейтинга нет, создаем пустой объект
            rating_stats = None

        # Получаем количество учеников
        try:
            students_count = teacher.student_set.count()
        except:
            students_count = 0

        teacher_data = {
            'id': teacher.id,
            'full_name': teacher.user.get_full_name(),
            'photo': teacher.user.photo,
            'subjects': teacher.subjects.all(),
            'experience': teacher.experience,
            'education': teacher.education,
            'bio': teacher.team_description or teacher.bio,
            'rating': rating_stats,
            'students_count': students_count,
            'email': teacher.user.email,
            'phone': teacher.user.phone,
            'social': {
                'telegram': getattr(teacher, 'telegram_url', ''),
                'whatsapp': getattr(teacher, 'whatsapp_url', ''),
                'vk': getattr(teacher, 'vk_url', ''),
                'instagram': getattr(teacher, 'instagram_url', ''),
            },
            'highlight': teacher.team_highlight,
            'sort_order': teacher.team_sort_order,
        }

        if teacher.team_highlight:
            highlighted_teachers.append(teacher_data)
        else:
            regular_teachers.append(teacher_data)

    context = {
        'highlighted_teachers': highlighted_teachers,
        'regular_teachers': regular_teachers,
        'page_title': 'Наша команда',
        'total_teachers': teachers.count(),
    }
    return render(request, 'school/team.html', context)

def teacher_detail(request, teacher_id):
    """
    Детальная страница учителя
    """
    teacher = get_object_or_404(
        Teacher.objects.select_related('user').prefetch_related(
            'subjects',
            'student_set__user'
        ),
        id=teacher_id
    )

    # Получаем публичные отзывы
    feedbacks = LessonFeedback.objects.filter(
        teacher=teacher,
        is_public=True
    ).select_related(
        'student__user',
        'lesson__subject'
    ).order_by('-created_at')[:20]

    # Безопасно получаем рейтинг
    try:
        rating_stats = teacher.rating_stats
    except:
        rating_stats = None

    context = {
        'teacher': teacher,
        'feedbacks': feedbacks,
        'rating_stats': rating_stats,
        'page_title': f'{teacher.user.get_full_name()} - учитель',
    }
    # ИСПРАВЛЕННЫЙ ПУТЬ
    return render(request, 'school/teacher/teacher_detail.html', context)


def home(request):
    """
    Главная страница
    """
    trial_form = TrialRequestForm()
    feedback_form = PublicFeedbackForm()

    # Получаем активные отзывы для главной
    feedbacks = Feedback.objects.filter(
        is_active=True,
        is_on_main=True
    ).select_related('teacher__user').order_by('sort_order', '-created_at')

    # ===== НОВАЯ СТАТИСТИКА ПО ОТЗЫВАМ =====
    # Все активные отзывы (не только на главной)
    all_active_feedbacks = Feedback.objects.filter(is_active=True)
    
    total_feedbacks = all_active_feedbacks.count()
    
    # Статистика по оценкам
    rating_stats = {
        'avg': all_active_feedbacks.aggregate(Avg('rating'))['rating__avg'] or 0,
        '5': all_active_feedbacks.filter(rating=5).count(),
        '4': all_active_feedbacks.filter(rating=4).count(),
        '3': all_active_feedbacks.filter(rating=3).count(),
        '2': all_active_feedbacks.filter(rating=2).count(),
        '1': all_active_feedbacks.filter(rating=1).count(),
    }
    
    # Проценты для визуализации
    if total_feedbacks > 0:
        rating_stats['5_percent'] = round(rating_stats['5'] / total_feedbacks * 100)
        rating_stats['4_percent'] = round(rating_stats['4'] / total_feedbacks * 100)
        rating_stats['3_percent'] = round(rating_stats['3'] / total_feedbacks * 100)
        rating_stats['2_percent'] = round(rating_stats['2'] / total_feedbacks * 100)
        rating_stats['1_percent'] = round(rating_stats['1'] / total_feedbacks * 100)
    else:
        rating_stats['5_percent'] = rating_stats['4_percent'] = rating_stats['3_percent'] = rating_stats['2_percent'] = rating_stats['1_percent'] = 0
    # ===== КОНЕЦ НОВОЙ СТАТИСТИКИ =====

    if request.method == 'POST':
        if 'trial_form' in request.POST:
            trial_form = TrialRequestForm(request.POST)
            if trial_form.is_valid():
                trial_form.save()
                messages.success(request, 'Заявка успешно отправлена! Мы свяжемся с вами в ближайшее время.')
                return redirect('home')
        elif 'feedback_form' in request.POST:
            feedback_form = PublicFeedbackForm(request.POST)
            if feedback_form.is_valid():
                feedback = feedback_form.save(commit=False)
                feedback.is_active = False
                feedback.is_on_main = False
                feedback.save()

                # Уведомление админу
                admin_users = User.objects.filter(is_superuser=True)
                for admin in admin_users:
                    Notification.objects.create(
                        user=admin,
                        title='✍️ Новый отзыв',
                        message=f'Новый отзыв от {feedback.name} ожидает модерации',
                        notification_type='system',
                        link=f'/admin/school/feedback/{feedback.id}/change/'
                    )

                messages.success(
                    request,
                    'Спасибо за ваш отзыв! Он появится на сайте после проверки.'
                )
                return redirect('home')

    context = {
        'trial_form': trial_form,
        'feedback_form': feedback_form,
        'subjects': Subject.objects.all(),
        'feedbacks': feedbacks,
        # ===== ДОБАВЛЯЕМ В КОНТЕКСТ =====
        'total_feedbacks': total_feedbacks,
        'rating_stats': rating_stats,
        # ===== КОНЕЦ ДОБАВЛЕНИЯ =====
    }
    return render(request, 'school/home.html', context)


def add_feedback(request):
    """
    Добавление отзыва посетителем сайта
    """
    if request.method == 'POST':
        form = PublicFeedbackForm(request.POST)
        if form.is_valid():
            feedback = form.save(commit=False)
            feedback.is_active = False  # Отзыв требует модерации
            feedback.is_on_main = False
            feedback.save()

            # Отправляем уведомление админу
            admin_users = User.objects.filter(is_superuser=True)
            for admin in admin_users:
                Notification.objects.create(
                    user=admin,
                    title='✍️ Новый отзыв',
                    message=f'Новый отзыв от {feedback.name} ожидает модерации',
                    notification_type='system',
                    link=f'/admin/school/feedback/{feedback.id}/change/'
                )

            messages.success(
                request,
                'Спасибо за ваш отзыв! Он появится на сайте после проверки модератором.'
            )
            return redirect('home')
        else:
            messages.error(request, 'Пожалуйста, исправьте ошибки в форме')
    else:
        form = PublicFeedbackForm()

    # Добавляем форму в контекст главной страницы
    context = {
        'feedback_form': form,
    }
    return render(request, 'school/add_feedback.html', context)